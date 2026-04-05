#!/usr/bin/env python3
"""app.py - JobAgent v2: AI-powered job board with resume matching"""

import io, csv, json, re, sqlite3, threading, time
from datetime import datetime
from pathlib import Path
from flask import Flask, Response, jsonify, request, send_file
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DB_PATH    = Path(__file__).parent / "jobs.db"
OLLAMA_API = "http://localhost:11434/api/generate"
OLLAMA_MODEL = "llama3.2"

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16 MB

# ── Database ──────────────────────────────────────────────────────────────────

def get_db():
    conn = sqlite3.connect(str(DB_PATH))
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db()
    conn.execute("""
        CREATE TABLE IF NOT EXISTS jobs (
            id             INTEGER PRIMARY KEY AUTOINCREMENT,
            title          TEXT,
            url            TEXT UNIQUE,
            published_date TEXT,
            source         TEXT,
            feed_name      TEXT,
            category       TEXT,
            description    TEXT,
            fetched_at     TEXT
        )
    """)
    conn.execute("CREATE INDEX IF NOT EXISTS idx_url      ON jobs(url)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_category ON jobs(category)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_source   ON jobs(source)")
    conn.execute("""
        CREATE TABLE IF NOT EXISTS saved_jobs (
            id       INTEGER PRIMARY KEY AUTOINCREMENT,
            job_id   INTEGER UNIQUE,
            saved_at TEXT DEFAULT (datetime('now'))
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS resume_sessions (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            parsed_json TEXT,
            created_at  TEXT DEFAULT (datetime('now'))
        )
    """)
    try:
        conn.execute("ALTER TABLE jobs ADD COLUMN location TEXT")
    except Exception:
        pass
    conn.commit()
    conn.close()

init_db()

# ── Location extraction ───────────────────────────────────────────────────────

def extract_location(title: str, description: str) -> str:
    t = f"{title or ''} {description or ''}".lower()
    if re.search(r"\b(remote|distributed|work.?from.?home|wfh)\b", t):
        if re.search(r"\b(us|usa|united states|america)\b.{0,40}\b(only|based|resident)\b|us-?only", t):
            return "Remote · US"
        if re.search(r"\b(uk|united kingdom|britain|england)\b", t):
            return "Remote · UK"
        if re.search(r"\beurope\b|\beu\b", t):
            return "Remote · Europe"
        if re.search(r"\bcanada\b", t):
            return "Remote · Canada"
        if re.search(r"\baustralia\b", t):
            return "Remote · APAC"
        return "Remote · Worldwide"
    for pat, label in [
        (r"san francisco|bay area|silicon valley", "San Francisco, US"),
        (r"new york|nyc",                          "New York, US"),
        (r"\blondon\b",                            "London, UK"),
        (r"\bberlin\b",                            "Berlin, DE"),
        (r"\btoronto\b|\bvancouver\b",             "Canada"),
        (r"\bamsterdam\b",                         "Amsterdam, NL"),
        (r"sydney|melbourne",                      "Australia"),
        (r"\bsingapore\b",                         "Singapore"),
        (r"\bparis\b",                             "Paris, FR"),
    ]:
        if re.search(pat, t):
            return label
    return "Remote · Worldwide"

# ── Resume parsing ────────────────────────────────────────────────────────────

SKILLS_DB = [
    "python","javascript","typescript","java","c++","c#","go","golang","rust","ruby","php","swift","kotlin","scala",
    "react","vue","angular","next.js","nuxt","svelte","node.js","express","django","flask","fastapi","spring","rails","laravel",
    "aws","azure","gcp","google cloud","docker","kubernetes","terraform","ansible","jenkins","github actions","ci/cd","linux","bash",
    "postgresql","mysql","mongodb","redis","elasticsearch","dynamodb","cassandra","snowflake","bigquery","supabase",
    "machine learning","deep learning","nlp","natural language processing","computer vision","pytorch","tensorflow","keras",
    "scikit-learn","xgboost","llm","large language model","openai","langchain","rag",
    "sql","nosql","graphql","rest","grpc","microservices","kafka","rabbitmq",
    "git","agile","scrum","jira","figma","sketch","product management","ux","ui design","user research",
    "data science","data analysis","data engineering","tableau","power bi","looker","spark","hadoop","airflow","dbt",
    "cybersecurity","penetration testing","soc","siem","devops","sre","cloud architecture","networking",
    "blockchain","solidity","web3","ios","android","react native","flutter",
]

def extract_text_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        return "\n".join(p.extract_text() or "" for p in reader.pages)
    except Exception:
        return ""

def extract_text_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception:
        return ""

def parse_simple(text: str) -> dict:
    t = text.lower()
    skills = [s for s in SKILLS_DB if s in t]
    title_re = re.compile(
        r"(?:^|\n)\s*([A-Z][a-zA-Z]+(?: [A-Z][a-zA-Z]+){0,3}"
        r"(?:Engineer|Developer|Scientist|Analyst|Manager|Designer|Architect|Lead|Director|Specialist|Consultant))",
        re.MULTILINE,
    )
    titles = list(dict.fromkeys(m.group(1).strip() for m in title_re.finditer(text)))[:5]
    years_list = [int(m) for m in re.findall(r"(\d{1,2})\+?\s*years?", t)]
    years = max(years_list) if years_list else 0
    edu = []
    for m in re.finditer(
        r"(bachelor|master|phd|ph\.d|mba|b\.s|m\.s|b\.a|m\.a|b\.eng|m\.eng)\.?\s+(?:of\s+|in\s+)?([\w\s,]+?)(?:\n|$|at|from|,)",
        text[:3000], re.IGNORECASE,
    ):
        edu.append({"degree": m.group(1).title(), "field": m.group(2).strip().title(), "institution": ""})
    return {
        "skills": skills[:25],
        "job_titles": titles or ["Technology Professional"],
        "education": edu[:3],
        "years_experience": years,
        "industries": ["Technology"],
        "summary": (
            f"Professional with expertise in {', '.join(skills[:4])}" if skills
            else "Technology professional"
        ),
    }

def parse_with_llama(text: str) -> dict | None:
    try:
        import requests as req
        prompt = (
            "Extract resume information and respond with ONLY a valid JSON object — no preamble, no markdown, no extra text.\n"
            'Keys required: {"skills": [...], "job_titles": [...], "education": [{"degree":"","field":"","institution":""}], '
            '"years_experience": <int>, "industries": [...], "summary": "2-sentence professional summary"}\n\n'
            f"Resume:\n{text[:3500]}\n\nJSON:"
        )
        r = req.post(OLLAMA_API, json={"model": OLLAMA_MODEL, "prompt": prompt, "stream": False}, timeout=90)
        if r.status_code == 200:
            raw = r.json().get("response", "")
            m = re.search(r"\{[\s\S]+\}", raw)
            if m:
                return json.loads(m.group())
    except Exception:
        pass
    return None

def parse_resume(file_bytes: bytes, filename: str) -> tuple[dict, bool]:
    fname = filename.lower()
    if fname.endswith(".pdf"):
        text = extract_text_pdf(file_bytes)
    elif fname.endswith(".docx"):
        text = extract_text_docx(file_bytes)
    else:
        text = file_bytes.decode("utf-8", errors="ignore")
    result = parse_with_llama(text)
    if result:
        return result, True
    return parse_simple(text), False

# ── Job matching ──────────────────────────────────────────────────────────────

def score_job(job: dict, resume: dict) -> tuple[int, list[str]]:
    title = (job.get("title") or "").lower()
    desc  = (job.get("description") or "").lower()
    text  = f"{title} {desc}"
    skills     = [s.lower() for s in resume.get("skills", [])]
    job_titles = [t.lower() for t in resume.get("job_titles", [])]
    years      = resume.get("years_experience", 0)
    score, reasons = 0, []
    hits = [s for s in skills if s in text]
    score += min(50, len(hits) * 8)
    if hits:
        reasons.append(f"Skills: {', '.join(hits[:4])}")
    STOP = {"and","or","the","a","an","in","of","at","for","to","with",""}
    for prev in job_titles:
        words   = set(re.split(r"\W+", prev)) - STOP
        overlap = [w for w in words if w and w in title]
        if prev in title or len(overlap) >= 2:
            score += 35; reasons.append(f"Matches role: {prev.title()}"); break
        elif overlap:
            score += 15
    senior = {"senior","staff","lead","principal","architect","director","vp"}
    junior = {"junior","associate","entry","intern"}
    tw = set(title.split())
    if years >= 5 and tw & senior:
        score += 15; reasons.append("Matches your seniority level")
    elif years < 3 and tw & junior:
        score += 15
    elif 2 <= years <= 7 and not (tw & senior) and not (tw & junior):
        score += 8
    return min(100, score), reasons

# ── Helpers ───────────────────────────────────────────────────────────────────

def rows_to_jobs(rows) -> list[dict]:
    jobs = []
    for r in rows:
        j = dict(r)
        if not j.get("location"):
            j["location"] = extract_location(j.get("title",""), j.get("description",""))
        jobs.append(j)
    return jobs

# ── Background refresh ────────────────────────────────────────────────────────

_REFRESH = {
    "running": False, "total": 0, "done": 0,
    "added": 0, "failed": 0, "feed": "", "finished_at": None,
}

def _do_refresh():
    try:
        from fetch_jobs import load_feeds, fetch_feed, upsert_job
        feeds = load_feeds()
        _REFRESH.update({"total": len(feeds), "done": 0, "added": 0, "failed": 0, "feed": ""})
        conn  = get_db()
        for feed in feeds:
            _REFRESH["feed"] = f"{feed['feed_name']} · {feed['source']}"
            try:
                jobs = fetch_feed(feed)
                for job in jobs:
                    upsert_job(conn, job)
                conn.commit()
                _REFRESH["added"] += len(jobs)
            except Exception:
                _REFRESH["failed"] += 1
            _REFRESH["done"] += 1
            time.sleep(0.5)
        conn.close()
    except Exception as exc:
        print(f"Refresh error: {exc}")
    finally:
        _REFRESH["running"]     = False
        _REFRESH["finished_at"] = datetime.utcnow().isoformat()

# ── API routes ────────────────────────────────────────────────────────────────

@app.route("/api/jobs")
def api_jobs():
    if not DB_PATH.exists():
        return jsonify([])
    conn = get_db()
    rows = conn.execute("SELECT * FROM jobs ORDER BY published_date DESC, fetched_at DESC").fetchall()
    conn.close()
    return jsonify(rows_to_jobs(rows))

@app.route("/api/stats")
def api_stats():
    if not DB_PATH.exists():
        return jsonify({"total":0,"categories":{},"sources":{},"dates":[]})
    conn  = get_db()
    cats  = dict(conn.execute("SELECT category, COUNT(*) FROM jobs GROUP BY category").fetchall())
    srcs  = dict(conn.execute("SELECT source,   COUNT(*) FROM jobs GROUP BY source").fetchall())
    dates = conn.execute(
        "SELECT DATE(published_date) d, COUNT(*) c FROM jobs "
        "WHERE published_date >= date('now','-14 days') GROUP BY d ORDER BY d"
    ).fetchall()
    total = conn.execute("SELECT COUNT(*) FROM jobs").fetchone()[0]
    conn.close()
    return jsonify({"total": total, "categories": cats, "sources": srcs,
                    "dates": [{"date": r[0], "count": r[1]} for r in dates]})

@app.route("/api/upload-resume", methods=["POST"])
def upload_resume():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "Empty filename"}), 400
    parsed, used_llama = parse_resume(f.read(), f.filename)
    conn = get_db()
    cur  = conn.execute("INSERT INTO resume_sessions (parsed_json) VALUES (?)", [json.dumps(parsed)])
    sid  = cur.lastrowid
    conn.commit(); conn.close()
    return jsonify({"session_id": sid, "parsed": parsed, "used_llama": used_llama})

@app.route("/api/match/<int:sid>")
def api_match(sid):
    conn = get_db()
    row  = conn.execute("SELECT parsed_json FROM resume_sessions WHERE id=?", [sid]).fetchone()
    if not row:
        conn.close(); return jsonify({"error": "Session not found"}), 404
    resume = json.loads(row[0])
    jobs   = rows_to_jobs(conn.execute("SELECT * FROM jobs").fetchall())
    conn.close()
    scored = []
    for j in jobs:
        sc, reasons = score_job(j, resume)
        if sc > 10:
            scored.append({**j, "match_score": sc, "match_reasons": reasons})
    scored.sort(key=lambda x: x["match_score"], reverse=True)
    return jsonify({"matches": scored[:100], "total": len(scored)})

@app.route("/api/saved-jobs", methods=["GET"])
def get_saved():
    conn = get_db()
    rows = conn.execute(
        "SELECT j.*, s.saved_at FROM saved_jobs s "
        "JOIN jobs j ON s.job_id = j.id ORDER BY s.saved_at DESC"
    ).fetchall()
    conn.close()
    return jsonify(rows_to_jobs(rows))

@app.route("/api/saved-jobs", methods=["POST"])
def add_saved():
    data = request.get_json() or {}
    jid  = data.get("job_id")
    if not jid:
        return jsonify({"error": "job_id required"}), 400
    conn = get_db()
    try:
        conn.execute("INSERT INTO saved_jobs (job_id) VALUES (?)", [jid])
        conn.commit()
    except sqlite3.IntegrityError:
        pass
    conn.close()
    return jsonify({"ok": True})

@app.route("/api/saved-jobs/<int:job_id>", methods=["DELETE"])
def remove_saved(job_id):
    conn = get_db()
    conn.execute("DELETE FROM saved_jobs WHERE job_id=?", [job_id])
    conn.commit(); conn.close()
    return jsonify({"ok": True})

@app.route("/api/saved-jobs/download")
def download_saved():
    conn = get_db()
    rows = conn.execute(
        "SELECT j.title, j.url, j.source, j.category, j.location, j.published_date, j.description "
        "FROM saved_jobs s JOIN jobs j ON s.job_id = j.id ORDER BY s.saved_at DESC"
    ).fetchall()
    conn.close()
    fmt = request.args.get("format", "csv")
    if fmt == "json":
        return Response(json.dumps([dict(r) for r in rows], indent=2), mimetype="application/json",
                        headers={"Content-Disposition": "attachment;filename=saved_jobs.json"})
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["Title","URL","Source","Category","Location","Date","Description"])
    for r in rows: w.writerow(list(r))
    buf.seek(0)
    return Response(buf.getvalue(), mimetype="text/csv",
                    headers={"Content-Disposition": "attachment;filename=saved_jobs.csv"})

@app.route("/api/match/<int:sid>/download")
def download_matched(sid):
    conn = get_db()
    row  = conn.execute("SELECT parsed_json FROM resume_sessions WHERE id=?", [sid]).fetchone()
    if not row:
        conn.close(); return jsonify({"error": "not found"}), 404
    resume = json.loads(row[0])
    jobs   = rows_to_jobs(conn.execute("SELECT * FROM jobs").fetchall())
    conn.close()
    scored = []
    for j in jobs:
        sc, reasons = score_job(j, resume)
        if sc > 10:
            scored.append({**j, "match_score": sc, "match_reasons": "; ".join(reasons)})
    scored.sort(key=lambda x: x["match_score"], reverse=True)
    fmt = request.args.get("format", "csv")
    if fmt == "json":
        return Response(json.dumps(scored[:100], indent=2), mimetype="application/json",
                        headers={"Content-Disposition": "attachment;filename=matched_jobs.json"})
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["Match%","Title","URL","Source","Category","Location","Date","Why Matched","Description"])
    for j in scored[:100]:
        w.writerow([j["match_score"], j["title"], j["url"], j["source"],
                    j.get("category",""), j.get("location",""), j.get("published_date",""),
                    j.get("match_reasons",""), (j.get("description") or "")[:300]])
    buf.seek(0)
    return Response(buf.getvalue(), mimetype="text/csv",
                    headers={"Content-Disposition": "attachment;filename=matched_jobs.csv"})

@app.route("/api/refresh", methods=["POST"])
def api_refresh_start():
    if _REFRESH["running"]:
        return jsonify({"error": "Already running"}), 409
    _REFRESH["running"] = True
    threading.Thread(target=_do_refresh, daemon=True).start()
    return jsonify({"ok": True})

@app.route("/api/refresh/status")
def api_refresh_status():
    return jsonify(_REFRESH)

# ── HTML template (raw string — no Jinja2 processing) ────────────────────────
DASHBOARD_HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>JobAgent</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Inter:ital,wght@0,300;0,400;0,500;0,600;0,700;1,400&display=swap" rel="stylesheet">
<script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.0/dist/chart.umd.min.js"></script>
<style>
/* ── Variables ── */
:root {
  --bg:        #09090f;
  --surf:      #0f1118;
  --surf-2:    #161b28;
  --surf-3:    #1d2235;
  --border:    #1f2640;
  --border-2:  #2a3354;
  --accent:    #6366f1;
  --accent-2:  #818cf8;
  --accent-bg: rgba(99,102,241,.14);
  --green:     #10b981;
  --amber:     #f59e0b;
  --red:       #ef4444;
  --teal:      #14b8a6;
  --txt:       #e2e8f5;
  --txt-2:     #8b95b0;
  --txt-3:     #4a526e;
  --r:         10px;
  --r-lg:      14px;
  --shadow:    0 8px 32px rgba(0,0,0,.45);
  --trans:     .18s ease;
}

/* ── Reset ── */
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
html{font-size:15px;scroll-behavior:smooth}
body{font-family:'Inter',system-ui,sans-serif;background:var(--bg);color:var(--txt);line-height:1.6;min-height:100vh;overflow-x:hidden}
a{color:inherit;text-decoration:none}
button{cursor:pointer;border:none;background:none;font:inherit;color:inherit}
input,select{font:inherit;color:inherit}
::-webkit-scrollbar{width:5px;height:5px}
::-webkit-scrollbar-track{background:transparent}
::-webkit-scrollbar-thumb{background:var(--border-2);border-radius:3px}

/* ── Nav ── */
.nav{
  position:sticky;top:0;z-index:200;
  height:58px;
  background:rgba(9,9,15,.9);
  backdrop-filter:blur(16px);
  border-bottom:1px solid var(--border);
  display:flex;align-items:center;gap:6px;
  padding:0 20px;
}
.nav-brand{
  display:flex;align-items:center;gap:9px;
  font-weight:700;font-size:1.05rem;letter-spacing:-.025em;
  margin-right:16px;white-space:nowrap;
}
.nav-brand-dot{
  width:30px;height:30px;border-radius:8px;
  background:linear-gradient(135deg,#6366f1,#a78bfa);
  display:flex;align-items:center;justify-content:center;
  font-size:15px;flex-shrink:0;
}
.nav-tab{
  display:flex;align-items:center;gap:7px;
  padding:6px 14px;border-radius:8px;
  color:var(--txt-2);font-size:.86rem;font-weight:500;
  transition:all var(--trans);white-space:nowrap;
}
.nav-tab:hover{color:var(--txt);background:var(--surf-2)}
.nav-tab.active{color:var(--txt);background:var(--surf-2)}
.nav-tab svg{opacity:.7;flex-shrink:0}
.nav-tab.active svg{opacity:1}
.tab-badge{
  background:var(--accent);color:#fff;
  font-size:.68rem;font-weight:700;
  padding:1px 6px;border-radius:10px;min-width:18px;text-align:center;
}
.nav-spacer{flex:1}
.nav-stat{color:var(--txt-3);font-size:.8rem;margin-right:8px;white-space:nowrap}
.nav-stat b{color:var(--txt-2)}
.btn-icon{
  width:34px;height:34px;border-radius:8px;flex-shrink:0;
  display:flex;align-items:center;justify-content:center;
  background:var(--surf-2);border:1px solid var(--border);
  color:var(--txt-2);transition:all var(--trans);
}
.btn-icon:hover{color:var(--txt);border-color:var(--border-2)}

/* ── Layout shell ── */
.shell{display:flex;height:calc(100vh - 58px);overflow:hidden}

/* ── Sidebar ── */
.sidebar{
  width:252px;flex-shrink:0;
  background:var(--surf);
  border-right:1px solid var(--border);
  overflow-y:auto;
  padding:18px 14px;
  display:flex;flex-direction:column;gap:20px;
}
.sb-label{
  font-size:.68rem;font-weight:700;letter-spacing:.09em;
  text-transform:uppercase;color:var(--txt-3);
  margin-bottom:7px;
}
.sb-section{display:flex;flex-direction:column}

/* search */
.s-wrap{position:relative}
.s-wrap svg{position:absolute;left:9px;top:50%;transform:translateY(-50%);color:var(--txt-3);pointer-events:none}
.s-input{
  width:100%;background:var(--surf-2);border:1px solid var(--border);
  border-radius:var(--r);padding:8px 10px 8px 32px;
  font-size:.84rem;outline:none;transition:border-color var(--trans);
}
.s-input::placeholder{color:var(--txt-3)}
.s-input:focus{border-color:var(--accent)}

/* date pills */
.date-pills{display:flex;flex-wrap:wrap;gap:5px}
.date-pill{
  padding:4px 11px;font-size:.76rem;font-weight:500;
  border-radius:6px;border:1px solid var(--border);
  background:var(--surf-2);color:var(--txt-2);
  transition:all var(--trans);
}
.date-pill:hover{border-color:var(--border-2);color:var(--txt)}
.date-pill.active{background:var(--accent-bg);border-color:var(--accent);color:var(--accent-2)}

/* select */
.sb-select{
  width:100%;background:var(--surf-2);border:1px solid var(--border);
  border-radius:var(--r);padding:8px 28px 8px 10px;
  font-size:.84rem;outline:none;cursor:pointer;appearance:none;
  background-image:url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='11' height='11' fill='%234a526e' viewBox='0 0 16 16'%3E%3Cpath d='M7.247 11.14 2.451 5.658C1.885 5.013 2.345 4 3.204 4h9.592a1 1 0 0 1 .753 1.659l-4.796 5.48a1 1 0 0 1-1.506 0z'/%3E%3C/svg%3E");
  background-repeat:no-repeat;background-position:right 9px center;
  transition:border-color var(--trans);
}
.sb-select:focus{border-color:var(--accent)}

/* check items */
.ck-list{display:flex;flex-direction:column;gap:4px}
.ck-item{
  display:flex;align-items:center;gap:7px;
  padding:3px 5px;border-radius:6px;cursor:pointer;
  transition:background var(--trans);
}
.ck-item:hover{background:var(--surf-2)}
.ck-item input[type=checkbox]{display:none}
.ck-box{
  width:15px;height:15px;border-radius:4px;flex-shrink:0;
  border:1.5px solid var(--border-2);background:var(--surf-2);
  display:flex;align-items:center;justify-content:center;
  transition:all var(--trans);
}
.ck-item input:checked ~ .ck-box{background:var(--accent);border-color:var(--accent)}
.ck-item input:checked ~ .ck-box::after{
  content:'';display:block;
  width:8px;height:5px;
  border-left:2px solid #fff;border-bottom:2px solid #fff;
  transform:rotate(-45deg) translate(1px,-1px);
}
.ck-name{font-size:.81rem;color:var(--txt-2);flex:1;line-height:1.3}
.ck-cnt{
  font-size:.69rem;color:var(--txt-3);
  background:var(--surf-3);padding:1px 6px;border-radius:9px;
}

/* sidebar buttons */
.btn-sb{
  display:flex;align-items:center;justify-content:center;gap:7px;
  width:100%;padding:8px 12px;border-radius:var(--r);
  font-size:.81rem;font-weight:500;
  border:1px solid var(--border);background:var(--surf-2);
  color:var(--txt-2);transition:all var(--trans);
}
.btn-sb:hover{border-color:var(--border-2);color:var(--txt)}
.btn-charts{border-color:var(--accent);background:var(--accent-bg);color:var(--accent-2)}
.btn-charts:hover{background:rgba(99,102,241,.25)}

/* ── Main ── */
.main{flex:1;overflow-y:auto;overflow-x:hidden}

/* toolbar */
.toolbar{
  position:sticky;top:0;z-index:10;
  background:rgba(9,9,15,.85);backdrop-filter:blur(10px);
  border-bottom:1px solid var(--border);
  padding:12px 22px;
  display:flex;align-items:center;gap:12px;flex-wrap:wrap;
}
.tbar-count{font-size:.83rem;color:var(--txt-2)}
.tbar-count strong{color:var(--txt)}
.tbar-spacer{flex:1}
.geo-wrap{display:flex;align-items:center;gap:7px}
.geo-lbl{font-size:.78rem;color:var(--txt-3);white-space:nowrap}
.geo-sel{
  background:var(--surf-2);border:1px solid var(--border);
  border-radius:var(--r);padding:5px 26px 5px 9px;
  font-size:.78rem;outline:none;cursor:pointer;appearance:none;
  background-image:url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='10' height='10' fill='%234a526e' viewBox='0 0 16 16'%3E%3Cpath d='M7.247 11.14 2.451 5.658C1.885 5.013 2.345 4 3.204 4h9.592a1 1 0 0 1 .753 1.659l-4.796 5.48a1 1 0 0 1-1.506 0z'/%3E%3C/svg%3E");
  background-repeat:no-repeat;background-position:right 7px center;
  transition:border-color var(--trans);color:var(--txt-2);
}
.geo-sel:focus{border-color:var(--accent)}

/* job grid */
.job-grid{
  display:grid;
  grid-template-columns:repeat(auto-fill,minmax(310px,1fr));
  gap:14px;padding:18px 22px;
}

/* job card */
.job-card{
  background:var(--surf);border:1px solid var(--border);
  border-radius:var(--r-lg);padding:16px;
  display:flex;flex-direction:column;gap:11px;
  position:relative;overflow:hidden;
  transition:all var(--trans);
}
.job-card::after{
  content:'';position:absolute;top:0;left:0;right:0;height:2px;
  background:transparent;transition:background var(--trans);
}
.job-card:hover{
  border-color:var(--border-2);
  box-shadow:0 6px 28px rgba(0,0,0,.35);
  transform:translateY(-2px);
}
.job-card:hover::after{background:linear-gradient(90deg,var(--accent),#a78bfa)}

.jc-top{display:flex;justify-content:space-between;align-items:flex-start;gap:8px}
.jc-title{font-size:.93rem;font-weight:600;line-height:1.4;flex:1}
.jc-title a{color:var(--txt)}
.jc-title a:hover{color:var(--accent-2);transition:color var(--trans)}
.save-btn{
  width:28px;height:28px;border-radius:7px;flex-shrink:0;
  display:flex;align-items:center;justify-content:center;
  background:var(--surf-2);border:1px solid var(--border);
  color:var(--txt-3);transition:all var(--trans);
}
.save-btn:hover{border-color:var(--amber);color:var(--amber)}
.save-btn.saved{background:rgba(245,158,11,.14);border-color:var(--amber);color:var(--amber)}

.jc-meta{display:flex;flex-wrap:wrap;gap:5px}
.badge{
  font-size:.69rem;font-weight:600;letter-spacing:.02em;
  padding:2px 8px;border-radius:20px;white-space:nowrap;
}
.b-source{background:rgba(99,102,241,.14);color:#a5b4fc;border:1px solid rgba(99,102,241,.28)}
.b-loc{background:rgba(20,184,166,.1);color:#5eead4;border:1px solid rgba(20,184,166,.22)}

.jc-score{margin-bottom:2px}
.score-row{display:flex;align-items:center;gap:8px;margin-bottom:4px}
.score-pill{
  font-size:.75rem;font-weight:700;padding:2px 9px;border-radius:20px;
}
.sp-hi{background:rgba(16,185,129,.18);color:#6ee7b7;border:1px solid rgba(16,185,129,.35)}
.sp-md{background:rgba(245,158,11,.18);color:#fcd34d;border:1px solid rgba(245,158,11,.35)}
.sp-lo{background:rgba(99,102,241,.18);color:#a5b4fc;border:1px solid rgba(99,102,241,.35)}
.score-reasons{font-size:.69rem;color:var(--txt-3);flex:1;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}
.score-bar{height:3px;border-radius:2px;background:var(--surf-3);overflow:hidden}
.score-bar-fill{height:100%;border-radius:2px;transition:width .5s ease}

.jc-desc{
  font-size:.8rem;color:var(--txt-2);line-height:1.6;
  display:-webkit-box;-webkit-line-clamp:3;-webkit-box-orient:vertical;overflow:hidden;
}
.jc-footer{
  display:flex;align-items:center;justify-content:space-between;
  border-top:1px solid var(--border);padding-top:9px;margin-top:auto;
  font-size:.73rem;color:var(--txt-3);
}
.jc-footer a{
  color:var(--accent-2);font-size:.76rem;font-weight:500;
  padding:3px 10px;border-radius:6px;border:1px solid rgba(99,102,241,.3);
  transition:all var(--trans);
}
.jc-footer a:hover{background:var(--accent-bg)}

/* load more */
.load-wrap{padding:22px;display:flex;justify-content:center}
.btn-load{
  padding:9px 26px;border-radius:var(--r);
  border:1px solid var(--border-2);background:var(--surf-2);
  color:var(--txt-2);font-size:.84rem;font-weight:500;
  transition:all var(--trans);
}
.btn-load:hover{border-color:var(--accent);color:var(--accent-2)}

/* empty */
.empty{
  display:flex;flex-direction:column;align-items:center;
  padding:72px 24px;gap:12px;text-align:center;
}
.empty-ico{font-size:44px;opacity:.35}
.empty-title{font-size:1.05rem;font-weight:600;color:var(--txt)}
.empty-text{color:var(--txt-3);font-size:.85rem;max-width:300px;line-height:1.5}

/* ── AI Match view ── */
.match-page{padding:32px 36px;max-width:1160px;margin:0 auto}
.match-hero{text-align:center;padding:32px 0 28px}
.match-hero h2{font-size:1.75rem;font-weight:700;letter-spacing:-.03em;margin-bottom:7px}
.match-hero p{color:var(--txt-2);font-size:.93rem}
.grad-text{
  background:linear-gradient(135deg,var(--accent) 20%,#a78bfa);
  -webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;
}

.upload-zone{
  border:2px dashed var(--border-2);border-radius:var(--r-lg);
  padding:52px 32px;text-align:center;cursor:pointer;
  transition:all var(--trans);position:relative;background:var(--surf);
  max-width:560px;margin:0 auto;
}
.upload-zone:hover,.upload-zone.dragover{border-color:var(--accent);background:var(--accent-bg)}
.upload-zone input[type=file]{position:absolute;inset:0;opacity:0;cursor:pointer;width:100%;height:100%}
.upload-ico{font-size:38px;margin-bottom:12px;opacity:.65}
.upload-title{font-size:.98rem;font-weight:600;margin-bottom:5px}
.upload-hint{color:var(--txt-3);font-size:.81rem}
.upload-hint b{color:var(--accent-2);font-weight:500}

.model-pill{
  display:inline-flex;align-items:center;gap:6px;
  padding:5px 13px;border-radius:20px;font-size:.76rem;font-weight:500;
  margin-top:14px;
}
.mp-llama{background:rgba(99,102,241,.14);color:#a5b4fc;border:1px solid rgba(99,102,241,.28)}
.mp-kw{background:rgba(245,158,11,.14);color:#fcd34d;border:1px solid rgba(245,158,11,.28)}
.mp-dot{width:6px;height:6px;border-radius:50%;background:currentColor;animation:pulse 2s infinite}
@keyframes pulse{0%,100%{opacity:1}50%{opacity:.35}}

.parse-loader{display:flex;flex-direction:column;align-items:center;gap:14px;padding:52px}
.spinner{
  width:38px;height:38px;border-radius:50%;
  border:3px solid var(--border-2);border-top-color:var(--accent);
  animation:spin .75s linear infinite;
}
@keyframes spin{to{transform:rotate(360deg)}}
.parse-text{color:var(--txt-2);font-size:.88rem}

.match-results{display:grid;grid-template-columns:270px 1fr;gap:22px;margin-top:28px}
.profile-card{
  background:var(--surf);border:1px solid var(--border);
  border-radius:var(--r-lg);padding:18px;height:fit-content;
  position:sticky;top:18px;
}
.profile-card h3{font-size:.9rem;font-weight:600;margin-bottom:14px;color:var(--txt)}
.p-section{margin-bottom:14px}
.p-section:last-child{margin-bottom:0}
.p-label{
  font-size:.66rem;font-weight:700;text-transform:uppercase;
  letter-spacing:.08em;color:var(--txt-3);margin-bottom:6px;
}
.skill-chip{
  display:inline-block;margin:2px;
  padding:2px 8px;border-radius:20px;font-size:.72rem;font-weight:500;
  background:rgba(99,102,241,.12);color:#a5b4fc;border:1px solid rgba(99,102,241,.22);
}
.p-text{font-size:.81rem;color:var(--txt-2);margin-bottom:3px}
.p-text strong{color:var(--txt)}
.p-summary{font-size:.8rem;color:var(--txt-2);line-height:1.55;margin-bottom:12px;font-style:italic}

.match-header{display:flex;align-items:center;gap:10px;margin-bottom:14px;flex-wrap:wrap}
.match-header h3{font-size:.95rem;font-weight:600}
.match-sub{color:var(--txt-3);font-size:.82rem}
.dload-group{display:flex;gap:7px;margin-left:auto}
.btn-dl{
  display:flex;align-items:center;gap:5px;
  padding:6px 13px;border-radius:var(--r);
  background:var(--accent-bg);border:1px solid rgba(99,102,241,.35);
  color:var(--accent-2);font-size:.78rem;font-weight:500;
  transition:all var(--trans);
}
.btn-dl:hover{background:rgba(99,102,241,.28)}

/* ── Saved view ── */
.saved-page{padding:28px 36px;max-width:880px;margin:0 auto}
.saved-hdr{display:flex;align-items:center;gap:10px;margin-bottom:24px;flex-wrap:wrap}
.saved-hdr h2{font-size:1.3rem;font-weight:700}
.saved-badge{
  background:var(--accent-bg);border:1px solid rgba(99,102,241,.3);
  color:var(--accent-2);font-size:.75rem;font-weight:600;
  padding:2px 9px;border-radius:20px;
}
.saved-list{display:flex;flex-direction:column;gap:11px}
.saved-item{
  background:var(--surf);border:1px solid var(--border);
  border-radius:var(--r-lg);padding:15px 18px;
  display:flex;gap:14px;align-items:flex-start;
  transition:all var(--trans);
}
.saved-item:hover{border-color:var(--border-2)}
.si-info{flex:1;min-width:0}
.si-title{font-size:.92rem;font-weight:600;margin-bottom:5px}
.si-title a:hover{color:var(--accent-2);transition:color var(--trans)}
.si-meta{display:flex;flex-wrap:wrap;gap:5px;margin-bottom:7px}
.si-desc{font-size:.79rem;color:var(--txt-2);line-height:1.55;
  display:-webkit-box;-webkit-line-clamp:2;-webkit-box-orient:vertical;overflow:hidden}
.btn-rm{
  padding:5px 11px;border-radius:var(--r);flex-shrink:0;
  border:1px solid var(--border-2);background:var(--surf-2);
  color:var(--txt-3);font-size:.75rem;font-weight:500;
  transition:all var(--trans);white-space:nowrap;align-self:flex-start;
}
.btn-rm:hover{border-color:var(--red);color:var(--red)}

/* ── Charts modal ── */
.modal-bg{
  position:fixed;inset:0;z-index:500;
  background:rgba(0,0,0,.72);backdrop-filter:blur(6px);
  display:none;align-items:center;justify-content:center;padding:20px;
}
.modal-bg.open{display:flex}
.modal{
  background:var(--surf);border:1px solid var(--border-2);
  border-radius:var(--r-lg);width:100%;max-width:880px;
  max-height:90vh;overflow-y:auto;
  box-shadow:var(--shadow);padding:26px;
}
.modal-hdr{display:flex;align-items:center;justify-content:space-between;margin-bottom:24px}
.modal-hdr h2{font-size:1.1rem;font-weight:700}
.btn-close{
  width:30px;height:30px;border-radius:8px;
  display:flex;align-items:center;justify-content:center;
  background:var(--surf-2);border:1px solid var(--border);
  color:var(--txt-2);transition:all var(--trans);font-size:1rem;
}
.btn-close:hover{color:var(--txt);border-color:var(--border-2)}
.charts-grid{display:grid;grid-template-columns:1fr 1fr;gap:18px}
.chart-card{
  background:var(--surf-2);border:1px solid var(--border);
  border-radius:var(--r);padding:18px;
}
.chart-lbl{font-size:.78rem;font-weight:600;color:var(--txt-2);margin-bottom:14px}
.chart-wrap{height:195px;position:relative}

/* ── Toast ── */
.toasts{
  position:fixed;bottom:22px;right:22px;z-index:9999;
  display:flex;flex-direction:column;gap:7px;pointer-events:none;
}
.toast{
  padding:11px 15px;border-radius:var(--r);
  background:var(--surf-3);border:1px solid var(--border-2);
  color:var(--txt);font-size:.83rem;font-weight:500;
  box-shadow:var(--shadow);
  display:flex;align-items:center;gap:8px;pointer-events:auto;
  transform:translateX(120%);transition:transform .28s cubic-bezier(.22,.68,0,1.18);
}
.toast.show{transform:translateX(0)}
.toast-ok{border-color:rgba(16,185,129,.4)}
.toast-err{border-color:rgba(239,68,68,.4)}

/* ── Category badges ── */
.b-cat{border:1px solid}
.bc-Engineering{background:rgba(16,185,129,.12);color:#6ee7b7;border-color:rgba(16,185,129,.28)}
.bc-AIML,.bc-DataScience,.bc-Data{background:rgba(139,92,246,.12);color:#c4b5fd;border-color:rgba(139,92,246,.28)}
.bc-DevOps{background:rgba(245,158,11,.12);color:#fcd34d;border-color:rgba(245,158,11,.28)}
.bc-Design{background:rgba(236,72,153,.12);color:#f9a8d4;border-color:rgba(236,72,153,.28)}
.bc-Security{background:rgba(239,68,68,.12);color:#fca5a5;border-color:rgba(239,68,68,.28)}
.bc-Product{background:rgba(14,165,233,.12);color:#7dd3fc;border-color:rgba(14,165,233,.28)}
.bc-General,.bc-Marketing,.bc-Other{background:rgba(100,116,139,.12);color:#94a3b8;border-color:rgba(100,116,139,.28)}

/* ── Responsive ── */
@media(max-width:860px){.sidebar{display:none}.match-results{grid-template-columns:1fr}.charts-grid{grid-template-columns:1fr}}
@media(max-width:580px){.job-grid{padding:12px;grid-template-columns:1fr}.match-page,.saved-page{padding:16px}.nav-stat{display:none}}

/* ── Expandable card ── */
.job-card{cursor:pointer}
.job-card .jc-desc{
  display:-webkit-box;-webkit-line-clamp:3;-webkit-box-orient:vertical;overflow:hidden;
  transition:none;
}
.job-card.expanded .jc-desc{
  display:block;overflow:visible;-webkit-line-clamp:unset;
}
.jc-extra{
  display:none;flex-direction:column;gap:10px;
  border-top:1px solid var(--border);padding-top:12px;
}
.job-card.expanded .jc-extra{display:flex}
.jc-salary{
  display:flex;align-items:center;gap:7px;
  font-size:.82rem;font-weight:600;color:var(--green);
}
.jc-skills{display:flex;flex-wrap:wrap;gap:5px}
.jc-skill-chip{
  font-size:.69rem;padding:2px 8px;border-radius:20px;font-weight:500;
  background:rgba(99,102,241,.12);color:#a5b4fc;border:1px solid rgba(99,102,241,.2);
}
.expand-hint{
  color:var(--txt-3);font-size:.72rem;font-weight:500;
  display:flex;align-items:center;gap:4px;transition:color var(--trans);margin-left:auto;
}
.job-card:hover .expand-hint{color:var(--txt-2)}
.expand-arrow{transition:transform var(--trans)}
.job-card.expanded .expand-arrow{transform:rotate(180deg)}

/* ── Refresh progress bar ── */
.refresh-bar{
  position:fixed;top:58px;left:0;right:0;height:2px;z-index:199;
  background:var(--border);display:none;overflow:hidden;
}
.refresh-bar.active{display:block}
.refresh-fill{
  height:100%;
  background:linear-gradient(90deg,var(--accent),#a78bfa,var(--teal));
  background-size:200% 100%;
  transition:width .6s ease;
  animation:shimmer 2s linear infinite;
}
@keyframes shimmer{0%{background-position:200% 0}100%{background-position:-200% 0}}

/* refresh button spin */
.btn-icon.spinning svg{animation:spin .7s linear infinite}
</style>
</head>
<body>

<!-- Nav -->
<nav class="nav">
  <div class="nav-brand">
    <div class="nav-brand-dot">🎯</div>
    JobAgent
  </div>

  <button class="nav-tab active" data-view="browse" onclick="switchView('browse')">
    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><rect x="3" y="3" width="7" height="7"/><rect x="14" y="3" width="7" height="7"/><rect x="14" y="14" width="7" height="7"/><rect x="3" y="14" width="7" height="7"/></svg>
    Browse
  </button>
  <button class="nav-tab" data-view="match" onclick="switchView('match')">
    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M12 2 2 7l10 5 10-5-10-5z"/><path d="M2 17l10 5 10-5"/><path d="M2 12l10 5 10-5"/></svg>
    AI Match
  </button>
  <button class="nav-tab" data-view="saved" onclick="switchView('saved')">
    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M19 21l-7-5-7 5V5a2 2 0 0 1 2-2h10a2 2 0 0 1 2 2z"/></svg>
    Saved <span class="tab-badge" id="tab-badge">0</span>
  </button>

  <a href="/profile" class="nav-tab" style="text-decoration:none">
    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M20 21v-2a4 4 0 0 0-4-4H8a4 4 0 0 0-4 4v2"/><circle cx="12" cy="7" r="4"/></svg>
    My Profile
  </a>

  <div class="nav-spacer"></div>
  <span class="nav-stat" id="nav-stat"><b>–</b> jobs</span>
  <button class="btn-icon" id="refresh-btn" onclick="startRefresh()" title="Refresh job listings">
    <svg width="15" height="15" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.2"><polyline points="23 4 23 10 17 10"/><polyline points="1 20 1 14 7 14"/><path d="M3.51 9a9 9 0 0 1 14.85-3.36L23 10M1 14l4.64 4.36A9 9 0 0 0 20.49 15"/></svg>
  </button>
  <button class="btn-icon" onclick="openCharts()" title="Charts &amp; Insights">
    <svg width="15" height="15" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><line x1="18" y1="20" x2="18" y2="10"/><line x1="12" y1="20" x2="12" y2="4"/><line x1="6" y1="20" x2="6" y2="14"/></svg>
  </button>
</nav>

<!-- Refresh progress bar -->
<div class="refresh-bar" id="refresh-bar">
  <div class="refresh-fill" id="refresh-fill" style="width:0%"></div>
</div>

<!-- Shell -->
<div class="shell">

  <!-- Sidebar -->
  <aside class="sidebar" id="sidebar">
    <div class="sb-section">
      <div class="sb-label">Search</div>
      <div class="s-wrap">
        <svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><circle cx="11" cy="11" r="8"/><line x1="21" y1="21" x2="16.65" y2="16.65"/></svg>
        <input class="s-input" id="s-input" type="text" placeholder="Title, skills, keywords…" oninput="debounceFilt()">
      </div>
    </div>

    <div class="sb-section">
      <div class="sb-label">Date Range</div>
      <div class="date-pills">
        <button class="date-pill" data-days="1" onclick="setDate(this)">Today</button>
        <button class="date-pill active" data-days="7" onclick="setDate(this)">7 days</button>
        <button class="date-pill" data-days="30" onclick="setDate(this)">30 days</button>
        <button class="date-pill" data-days="0" onclick="setDate(this)">All time</button>
      </div>
    </div>

    <div class="sb-section">
      <div class="sb-label">Sort By</div>
      <select class="sb-select" id="sort-sel" onchange="applyFilt()">
        <option value="newest">Newest first</option>
        <option value="oldest">Oldest first</option>
        <option value="source">Source A–Z</option>
        <option value="title">Title A–Z</option>
      </select>
    </div>

    <div class="sb-section">
      <div class="sb-label">Category</div>
      <div class="ck-list" id="cat-list"></div>
    </div>

    <div class="sb-section">
      <div class="sb-label">Source</div>
      <div class="ck-list" id="src-list"></div>
    </div>

    <div class="sb-section" style="gap:8px;display:flex;flex-direction:column">
      <button class="btn-sb" onclick="clearFilt()">
        <svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><polyline points="1 4 1 10 7 10"/><path d="M3.51 15a9 9 0 1 0 .49-3.43"/></svg>
        Reset Filters
      </button>
      <button class="btn-sb btn-charts" onclick="openCharts()">
        <svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><line x1="18" y1="20" x2="18" y2="10"/><line x1="12" y1="20" x2="12" y2="4"/><line x1="6" y1="20" x2="6" y2="14"/></svg>
        Visualize Data
      </button>
    </div>
  </aside>

  <!-- Main area -->
  <main class="main" id="main">

    <!-- ── Browse View ── -->
    <div id="v-browse">
      <div class="toolbar">
        <span class="tbar-count">Showing <strong id="cnt-vis">0</strong> of <strong id="cnt-tot">0</strong> jobs</span>
        <div class="tbar-spacer"></div>
        <div class="geo-wrap">
          <span class="geo-lbl">Region:</span>
          <select class="geo-sel" id="geo-sel" onchange="applyFilt()">
            <option value="">All regions</option>
          </select>
        </div>
      </div>
      <div class="job-grid" id="job-grid"></div>
      <div class="load-wrap" id="load-wrap" style="display:none">
        <button class="btn-load" onclick="loadMore()">Load more jobs</button>
      </div>
    </div>

    <!-- ── AI Match View ── -->
    <div id="v-match" style="display:none">
      <div class="match-page">
        <div class="match-hero">
          <h2>AI Job <span class="grad-text">Matcher</span></h2>
          <p>Upload your resume — we'll parse your skills and surface your best opportunities</p>
        </div>

        <!-- Upload state -->
        <div id="upload-state">
          <div class="upload-zone" id="drop-zone">
            <input type="file" id="resume-file" accept=".pdf,.docx,.txt" onchange="onFilePick(event)">
            <div class="upload-ico">📄</div>
            <div class="upload-title">Drop your resume here</div>
            <div class="upload-hint">or <b>click to browse</b> &nbsp;·&nbsp; PDF, DOCX or TXT</div>
          </div>
          <div style="text-align:center">
            <span class="model-pill mp-kw" id="mode-hint">
              <span class="mp-dot"></span>
              Keyword mode — <a href="https://ollama.com" target="_blank" style="color:inherit;text-decoration:underline">install Ollama</a> + <code style="font-size:.7rem">ollama pull llama3.2</code> for AI mode
            </span>
          </div>
        </div>

        <!-- Parsing state -->
        <div id="parse-state" style="display:none">
          <div class="parse-loader">
            <div class="spinner"></div>
            <div class="parse-text" id="parse-msg">Parsing resume…</div>
          </div>
        </div>

        <!-- Results state -->
        <div id="result-state" style="display:none">
          <div style="text-align:center;margin-bottom:20px">
            <span class="model-pill" id="used-model"></span>
          </div>
          <div class="match-results">
            <div class="profile-card" id="profile-card"></div>
            <div>
              <div class="match-header">
                <h3>Top Matches</h3>
                <span class="match-sub" id="match-sub"></span>
                <div class="dload-group">
                  <button class="btn-dl" onclick="dlMatch('csv')">
                    <svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/><polyline points="7 10 12 15 17 10"/><line x1="12" y1="15" x2="12" y2="3"/></svg>
                    CSV
                  </button>
                  <button class="btn-dl" onclick="dlMatch('json')">
                    <svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/><polyline points="7 10 12 15 17 10"/><line x1="12" y1="15" x2="12" y2="3"/></svg>
                    JSON
                  </button>
                </div>
              </div>
              <div class="job-grid" id="match-grid" style="padding:0;grid-template-columns:repeat(auto-fill,minmax(290px,1fr))"></div>
            </div>
          </div>
          <div style="text-align:center;margin-top:20px">
            <button class="btn-sb" style="width:auto;display:inline-flex;padding:8px 20px" onclick="resetMatch()">
              <svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><polyline points="1 4 1 10 7 10"/><path d="M3.51 15a9 9 0 1 0 .49-3.43"/></svg>
              Upload different resume
            </button>
          </div>
        </div>
      </div>
    </div>

    <!-- ── Saved View ── -->
    <div id="v-saved" style="display:none">
      <div class="saved-page">
        <div class="saved-hdr">
          <h2>Saved Jobs</h2>
          <span class="saved-badge" id="saved-badge">0 jobs</span>
          <div class="dload-group" style="margin-left:auto">
            <button class="btn-dl" onclick="dlSaved('csv')">
              <svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/><polyline points="7 10 12 15 17 10"/><line x1="12" y1="15" x2="12" y2="3"/></svg>
              Download CSV
            </button>
            <button class="btn-dl" onclick="dlSaved('json')" style="background:rgba(20,184,166,.1);border-color:rgba(20,184,166,.3);color:#5eead4">
              <svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/><polyline points="7 10 12 15 17 10"/><line x1="12" y1="15" x2="12" y2="3"/></svg>
              Download JSON
            </button>
          </div>
        </div>
        <div class="saved-list" id="saved-list"></div>
      </div>
    </div>

  </main>
</div>

<!-- Charts Modal -->
<div class="modal-bg" id="charts-modal" onclick="if(event.target===this)closeCharts()">
  <div class="modal">
    <div class="modal-hdr">
      <h2>📊 Job Market Insights</h2>
      <button class="btn-close" onclick="closeCharts()">✕</button>
    </div>
    <div class="charts-grid">
      <div class="chart-card">
        <div class="chart-lbl">Jobs by Category</div>
        <div class="chart-wrap"><canvas id="chart-cat"></canvas></div>
      </div>
      <div class="chart-card">
        <div class="chart-lbl">Jobs by Source</div>
        <div class="chart-wrap"><canvas id="chart-src"></canvas></div>
      </div>
      <div class="chart-card">
        <div class="chart-lbl">Region Distribution</div>
        <div class="chart-wrap"><canvas id="chart-geo"></canvas></div>
      </div>
      <div class="chart-card">
        <div class="chart-lbl">Posting Activity — Last 14 Days</div>
        <div class="chart-wrap"><canvas id="chart-time"></canvas></div>
      </div>
    </div>
  </div>
</div>

<!-- Toasts -->
<div class="toasts" id="toasts"></div>

<script>
"use strict";

// ── State ────────────────────────────────────────────────────────────────────
let ALL_JOBS    = [];
let FILTERED    = [];
let SAVED_IDS   = new Set();
let PAGE        = 0;
const PG_SIZE   = 60;
let filterDays  = 7;
let filtTimer   = null;
let sessionId   = null;
let charts      = {};

// ── Boot ─────────────────────────────────────────────────────────────────────
async function boot() {
  await Promise.all([loadJobs(), loadSaved()]);
  buildCheckLists();
  applyFilt();
  checkOllama();
}

async function loadJobs() {
  try {
    const r = await fetch('/api/jobs');
    ALL_JOBS = await r.json();
    document.getElementById('nav-stat').innerHTML = `<b>${ALL_JOBS.length.toLocaleString()}</b> jobs`;
    // Populate geo dropdown
    const locs = [...new Set(ALL_JOBS.map(j => j.location).filter(Boolean))].sort();
    const sel  = document.getElementById('geo-sel');
    locs.forEach(l => { const o = document.createElement('option'); o.value = l; o.textContent = l; sel.appendChild(o); });
  } catch(e) { toast('Could not load jobs', '⚠️', 'toast-err'); }
}

async function loadSaved() {
  try {
    const r    = await fetch('/api/saved-jobs');
    const data = await r.json();
    SAVED_IDS  = new Set(data.map(j => j.id));
    updateCounts();
    renderSaved(data);
  } catch(e) {}
}

async function checkOllama() {
  try {
    const r = await fetch('http://localhost:11434/api/tags', { signal: AbortSignal.timeout(1500) });
    if (r.ok) {
      document.getElementById('mode-hint').className = 'model-pill mp-llama';
      document.getElementById('mode-hint').innerHTML = '<span class="mp-dot"></span> Ollama detected — Llama AI mode active';
    }
  } catch(e) {}
}

// ── View switching ────────────────────────────────────────────────────────────
function switchView(v) {
  ['browse','match','saved'].forEach(name => {
    document.getElementById(`v-${name}`).style.display = v === name ? 'block' : 'none';
  });
  document.getElementById('sidebar').style.display = v === 'browse' ? 'flex' : 'none';
  document.querySelectorAll('.nav-tab').forEach(t => t.classList.toggle('active', t.dataset.view === v));
  if (v === 'saved') refreshSaved();
}

// ── Filter helpers ────────────────────────────────────────────────────────────
function buildCheckLists() {
  const cats = {}, srcs = {};
  ALL_JOBS.forEach(j => {
    if (j.category) cats[j.category] = (cats[j.category]||0)+1;
    if (j.source)   srcs[j.source]   = (srcs[j.source]||0)+1;
  });
  buildList('cat-list', cats, 'c-');
  buildList('src-list', srcs, 's-');
}
function buildList(elId, counts, pfx) {
  const el = document.getElementById(elId);
  el.innerHTML = Object.entries(counts).sort(([,a],[,b])=>b-a).map(([name,cnt])=>`
    <label class="ck-item">
      <input type="checkbox" id="${pfx}${name}" checked onchange="applyFilt()">
      <span class="ck-box"></span>
      <span class="ck-name">${esc(name)}</span>
      <span class="ck-cnt">${cnt}</span>
    </label>`).join('');
}
function setDate(btn) {
  document.querySelectorAll('.date-pill').forEach(b=>b.classList.remove('active'));
  btn.classList.add('active');
  filterDays = parseInt(btn.dataset.days);
  applyFilt();
}
function debounceFilt() {
  clearTimeout(filtTimer);
  filtTimer = setTimeout(applyFilt, 180);
}
function applyFilt() {
  const q    = (document.getElementById('s-input').value||'').toLowerCase();
  const sort = document.getElementById('sort-sel').value;
  const geo  = document.getElementById('geo-sel').value;
  const cats = new Set([...document.querySelectorAll('[id^="c-"]:checked')].map(c=>c.id.slice(2)));
  const srcs = new Set([...document.querySelectorAll('[id^="s-"]:checked')].map(c=>c.id.slice(2)));
  const cut  = filterDays > 0 ? new Date(Date.now() - filterDays * 864e5) : null;

  FILTERED = ALL_JOBS.filter(j => {
    if (!cats.has(j.category)) return false;
    if (!srcs.has(j.source))   return false;
    if (geo  && j.location !== geo) return false;
    if (cut  && j.published_date && new Date(j.published_date) < cut) return false;
    if (q) {
      const hay = `${j.title} ${j.source} ${j.feed_name} ${j.description} ${j.category} ${j.location}`.toLowerCase();
      if (!hay.includes(q)) return false;
    }
    return true;
  });

  FILTERED.sort((a,b) => {
    if (sort==='newest') return (b.published_date||'').localeCompare(a.published_date||'');
    if (sort==='oldest') return (a.published_date||'').localeCompare(b.published_date||'');
    if (sort==='source') return (a.source||'').localeCompare(b.source||'');
    if (sort==='title')  return (a.title||'').localeCompare(b.title||'');
    return 0;
  });

  PAGE = 0;
  renderJobs();
}
function clearFilt() {
  document.getElementById('s-input').value = '';
  document.getElementById('sort-sel').value = 'newest';
  document.getElementById('geo-sel').value  = '';
  document.querySelectorAll('.date-pill').forEach(b=>b.classList.toggle('active', b.dataset.days==='7'));
  filterDays = 7;
  document.querySelectorAll('[id^="c-"],[id^="s-"]').forEach(c=>c.checked=true);
  applyFilt();
}
function loadMore() { PAGE++; renderJobs(); }

// ── Render jobs ───────────────────────────────────────────────────────────────
function renderJobs() {
  const slice = FILTERED.slice(0, (PAGE+1)*PG_SIZE);
  document.getElementById('cnt-vis').textContent = slice.length.toLocaleString();
  document.getElementById('cnt-tot').textContent = FILTERED.length.toLocaleString();
  const grid = document.getElementById('job-grid');
  grid.innerHTML = slice.length
    ? slice.map(j => cardHTML(j, false)).join('')
    : `<div class="empty" style="grid-column:1/-1">
         <div class="empty-ico">🔍</div>
         <div class="empty-title">No jobs found</div>
         <div class="empty-text">Try adjusting your search or filters</div>
       </div>`;
  document.getElementById('load-wrap').style.display = FILTERED.length > slice.length ? 'flex' : 'none';
}

// ── Salary + skill extraction ─────────────────────────────────────────────────
const SKILL_WORDS = [
  'python','javascript','typescript','java','go','golang','rust','c++','c#','ruby','php','swift','kotlin','scala',
  'react','vue','angular','next.js','node.js','express','django','flask','fastapi','spring','rails',
  'aws','azure','gcp','docker','kubernetes','terraform','linux','bash','ci/cd','github actions',
  'postgresql','mysql','mongodb','redis','elasticsearch','dynamodb','kafka','rabbitmq',
  'machine learning','deep learning','nlp','pytorch','tensorflow','scikit-learn','llm','langchain',
  'sql','graphql','rest','grpc','microservices','spark','airflow','dbt','snowflake','bigquery',
  'figma','sketch','ux','ui design','product management','agile','scrum',
  'cybersecurity','penetration testing','devops','sre','blockchain','solidity','web3',
  'ios','android','react native','flutter','data science','data engineering','tableau','power bi',
];

function extractSalary(text) {
  if (!text) return null;
  const m = text.match(/\$[\d,]+k?(?:\s*[-–]\s*\$[\d,]+k?)?(?:\s*\/?\s*(?:yr|year|annual|mo|month|hour|hr))?/i)
           || text.match(/(?:salary|compensation|pay|rate)[:\s]+[\$£€]?\d[\d,k.]+.{0,25}/i);
  if (!m) return null;
  return m[0].replace(/\s+/g,' ').trim().slice(0,60);
}

function extractSkills(text) {
  if (!text) return [];
  const t = text.toLowerCase();
  return SKILL_WORDS.filter(s => t.includes(s)).slice(0,12);
}

function cardHTML(j, showScore) {
  const saved    = SAVED_IDS.has(j.id);
  const catKey   = (j.category||'General').replace(/[^a-zA-Z]/g,'');
  const dateStr  = timeAgo(j.published_date);
  const fullDesc = j.description || '';
  const descSnip = esc(fullDesc.slice(0, 220));
  const loc      = j.location || '';
  const salary   = extractSalary(fullDesc);
  const skills   = extractSkills(fullDesc);

  let scoreHTML = '';
  if (showScore && j.match_score != null) {
    const sc    = j.match_score;
    const cls   = sc>=70 ? 'sp-hi' : sc>=40 ? 'sp-md' : 'sp-lo';
    const color = sc>=70 ? 'var(--green)' : sc>=40 ? 'var(--amber)' : 'var(--accent)';
    const why   = (j.match_reasons||[]).join(' · ');
    scoreHTML   = `<div class="jc-score">
      <div class="score-row">
        <span class="score-pill ${cls}">${sc}% match</span>
        <span class="score-reasons" title="${esc(why)}">${esc(why)}</span>
      </div>
      <div class="score-bar"><div class="score-bar-fill" style="width:${sc}%;background:${color}"></div></div>
    </div>`;
  }

  const extraHTML = `<div class="jc-extra">
    ${salary ? `<div class="jc-salary">💰 ${esc(salary)}</div>` : ''}
    ${skills.length ? `<div>
      <div style="font-size:.67rem;font-weight:700;text-transform:uppercase;letter-spacing:.07em;color:var(--txt-3);margin-bottom:5px">Skills mentioned</div>
      <div class="jc-skills">${skills.map(s=>`<span class="jc-skill-chip">${esc(s)}</span>`).join('')}</div>
    </div>` : ''}
    ${fullDesc.length > 220 ? `<div>
      <div style="font-size:.67rem;font-weight:700;text-transform:uppercase;letter-spacing:.07em;color:var(--txt-3);margin-bottom:5px">Full Description</div>
      <div style="font-size:.8rem;color:var(--txt-2);line-height:1.65;white-space:pre-wrap">${esc(fullDesc)}</div>
    </div>` : ''}
    ${!salary && !skills.length && fullDesc.length <= 220 ? `<div style="font-size:.8rem;color:var(--txt-3);font-style:italic">No additional details in the listing preview — click View to see the full posting.</div>` : ''}
  </div>`;

  return `<div class="job-card" data-id="${j.id}">
    <div class="jc-top">
      <div class="jc-title"><a href="${esc(j.url||'#')}" target="_blank" rel="noopener">${esc(j.title||'Untitled')}</a></div>
      <button class="save-btn${saved?' saved':''}" onclick="toggleSave(${j.id},this)" title="${saved?'Remove from saved':'Save job'}">
        <svg width="13" height="13" viewBox="0 0 24 24" fill="${saved?'currentColor':'none'}" stroke="currentColor" stroke-width="2"><path d="M19 21l-7-5-7 5V5a2 2 0 0 1 2-2h10a2 2 0 0 1 2 2z"/></svg>
      </button>
    </div>
    ${scoreHTML}
    <div class="jc-meta">
      <span class="badge b-source">${esc(j.source||'')}</span>
      ${j.category?`<span class="badge b-cat bc-${catKey}">${esc(j.category)}</span>`:''}
      ${loc?`<span class="badge b-loc">📍 ${esc(loc)}</span>`:''}
    </div>
    ${descSnip?`<div class="jc-desc">${descSnip}</div>`:''}
    ${extraHTML}
    <div class="jc-footer">
      <span>${dateStr}</span>
      <span class="expand-hint">
        <svg class="expand-arrow" width="11" height="11" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><polyline points="6 9 12 15 18 9"/></svg>
        Details
      </span>
      <a href="${esc(j.url||'#')}" target="_blank" rel="noopener">View →</a>
    </div>
  </div>`;
}

// ── Save / Unsave ─────────────────────────────────────────────────────────────
async function toggleSave(jobId, btn) {
  const wasSaved = SAVED_IDS.has(jobId);
  // Optimistic UI
  btn.classList.toggle('saved', !wasSaved);
  btn.querySelector('svg').setAttribute('fill', wasSaved ? 'none' : 'currentColor');
  if (wasSaved) { SAVED_IDS.delete(jobId); } else { SAVED_IDS.add(jobId); }
  updateCounts();
  try {
    if (wasSaved) {
      await fetch(`/api/saved-jobs/${jobId}`, {method:'DELETE'});
      toast('Removed from saved', '🗑️');
    } else {
      await fetch('/api/saved-jobs', {method:'POST', headers:{'Content-Type':'application/json'}, body:JSON.stringify({job_id:jobId})});
      toast('Job saved!', '🔖', 'toast-ok');
    }
  } catch(e) {
    // revert on error
    btn.classList.toggle('saved', wasSaved);
    btn.querySelector('svg').setAttribute('fill', wasSaved ? 'currentColor' : 'none');
    if (wasSaved) { SAVED_IDS.add(jobId); } else { SAVED_IDS.delete(jobId); }
    updateCounts();
    toast('Network error', '⚠️', 'toast-err');
  }
}

function updateCounts() {
  const n = SAVED_IDS.size;
  document.getElementById('tab-badge').textContent = n;
  document.getElementById('saved-badge').textContent = `${n} job${n!==1?'s':''}`;
}

async function refreshSaved() {
  try {
    const r    = await fetch('/api/saved-jobs');
    const data = await r.json();
    SAVED_IDS  = new Set(data.map(j => j.id));
    updateCounts();
    renderSaved(data);
  } catch(e) {}
}

function renderSaved(list) {
  const el = document.getElementById('saved-list');
  if (!list.length) {
    el.innerHTML = `<div class="empty">
      <div class="empty-ico">🔖</div>
      <div class="empty-title">No saved jobs yet</div>
      <div class="empty-text">Browse jobs and tap the bookmark on any card</div>
    </div>`;
    return;
  }
  el.innerHTML = list.map(j => {
    const catKey = (j.category||'General').replace(/[^a-zA-Z]/g,'');
    return `<div class="saved-item" data-id="${j.id}">
      <div class="si-info">
        <div class="si-title"><a href="${esc(j.url||'#')}" target="_blank" rel="noopener">${esc(j.title||'Untitled')}</a></div>
        <div class="si-meta">
          <span class="badge b-source">${esc(j.source||'')}</span>
          ${j.category?`<span class="badge b-cat bc-${catKey}">${esc(j.category)}</span>`:''}
          ${j.location?`<span class="badge b-loc">📍 ${esc(j.location)}</span>`:''}
        </div>
        <div class="si-desc">${esc((j.description||'').slice(0,200))}</div>
      </div>
      <button class="btn-rm" onclick="removeSavedItem(${j.id}, this.closest('.saved-item'))">Remove</button>
    </div>`;
  }).join('');
}

async function removeSavedItem(jobId, el) {
  try {
    await fetch(`/api/saved-jobs/${jobId}`, {method:'DELETE'});
    SAVED_IDS.delete(jobId);
    el.remove();
    updateCounts();
    if (!document.querySelectorAll('.saved-item').length) renderSaved([]);
    toast('Removed', '🗑️');
    // Update bookmark buttons in browse view
    document.querySelectorAll(`.job-card[data-id="${jobId}"] .save-btn`).forEach(b => {
      b.classList.remove('saved');
      b.querySelector('svg').setAttribute('fill', 'none');
    });
  } catch(e) { toast('Error', '⚠️', 'toast-err'); }
}

function dlSaved(fmt) { location.href = `/api/saved-jobs/download?format=${fmt}`; }

// ── Resume / AI Match ─────────────────────────────────────────────────────────
const dropZone = document.getElementById('drop-zone');
dropZone.addEventListener('dragover',  e => { e.preventDefault(); dropZone.classList.add('dragover'); });
dropZone.addEventListener('dragleave', ()=> dropZone.classList.remove('dragover'));
dropZone.addEventListener('drop', e => {
  e.preventDefault(); dropZone.classList.remove('dragover');
  const file = e.dataTransfer.files[0];
  if (file) processResume(file);
});

function onFilePick(e) {
  const file = e.target.files[0];
  if (file) processResume(file);
}

async function processResume(file) {
  show('parse-state'); hide('upload-state'); hide('result-state');
  document.getElementById('parse-msg').textContent = `Reading ${file.name}…`;

  const fd = new FormData();
  fd.append('file', file);

  try {
    document.getElementById('parse-msg').textContent = 'Analyzing with AI…';
    const res  = await fetch('/api/upload-resume', {method:'POST', body:fd});
    const data = await res.json();
    if (data.error) throw new Error(data.error);

    sessionId = data.session_id;
    document.getElementById('parse-msg').textContent = 'Finding best matches…';

    const mres    = await fetch(`/api/match/${sessionId}`);
    const mdata   = await mres.json();

    hide('parse-state'); show('result-state');

    const pill = document.getElementById('used-model');
    if (data.used_llama) {
      pill.className = 'model-pill mp-llama';
      pill.innerHTML = '<span class="mp-dot"></span> Analyzed with Llama 3.2';
    } else {
      pill.className = 'model-pill mp-kw';
      pill.innerHTML = '<span class="mp-dot"></span> Keyword matching mode';
    }

    renderProfile(data.parsed);
    renderMatches(mdata.matches, mdata.total);
  } catch(e) {
    hide('parse-state'); show('upload-state');
    toast(`Error: ${e.message}`, '❌', 'toast-err');
  }
}

function renderProfile(p) {
  const skills = (p.skills||[]).slice(0,22);
  const edu    = (p.education||[]);
  document.getElementById('profile-card').innerHTML = `
    <h3>Your Profile</h3>
    ${p.summary ? `<p class="p-summary">${esc(p.summary)}</p>` : ''}
    <div class="p-section">
      <div class="p-label">Skills</div>
      <div>${skills.map(s=>`<span class="skill-chip">${esc(s)}</span>`).join('') || '<span style="color:var(--txt-3);font-size:.8rem">None detected</span>'}</div>
    </div>
    <div class="p-section">
      <div class="p-label">Experience</div>
      ${(p.job_titles||[]).map(t=>`<div class="p-text"><strong>${esc(t)}</strong></div>`).join('')}
      ${p.years_experience ? `<div class="p-text">${p.years_experience}+ years</div>` : ''}
    </div>
    ${edu.length ? `<div class="p-section">
      <div class="p-label">Education</div>
      ${edu.map(e=>`<div class="p-text">${esc(e.degree||'')}${e.field?' in '+esc(e.field):''}${e.institution?' · '+esc(e.institution):''}</div>`).join('')}
    </div>` : ''}
    ${(p.industries||[]).length ? `<div class="p-section">
      <div class="p-label">Industries</div>
      <div>${p.industries.map(i=>`<span class="skill-chip">${esc(i)}</span>`).join('')}</div>
    </div>` : ''}
  `;
}

function renderMatches(matches, total) {
  document.getElementById('match-sub').textContent = `${total} matched`;
  const grid = document.getElementById('match-grid');
  grid.innerHTML = matches.length
    ? matches.map(j => cardHTML(j, true)).join('')
    : `<div class="empty" style="grid-column:1/-1">
         <div class="empty-ico">🎯</div>
         <div class="empty-title">No strong matches</div>
         <div class="empty-text">Add more skills or job titles to your resume</div>
       </div>`;
}

function dlMatch(fmt) { if (sessionId) location.href = `/api/match/${sessionId}/download?format=${fmt}`; }

function resetMatch() {
  sessionId = null;
  document.getElementById('resume-file').value = '';
  show('upload-state'); hide('parse-state'); hide('result-state');
}

// ── Charts ────────────────────────────────────────────────────────────────────
const COLORS = ['#6366f1','#a78bfa','#10b981','#f59e0b','#ef4444','#06b6d4','#ec4899','#84cc16','#f97316','#8b5cf6'];

async function openCharts() {
  document.getElementById('charts-modal').classList.add('open');
  Object.values(charts).forEach(c => c.destroy());
  charts = {};
  try {
    const r     = await fetch('/api/stats');
    const stats = await r.json();
    const ticks = { color:'#4a526e', font:{size:10} };
    const grid  = { color:'#1f2640' };

    // Category donut
    const catK = Object.keys(stats.categories);
    charts.cat = new Chart(document.getElementById('chart-cat'), {
      type: 'doughnut',
      data: { labels: catK, datasets:[{ data: catK.map(k=>stats.categories[k]), backgroundColor: COLORS, borderWidth:0, hoverOffset:4 }] },
      options: { cutout:'58%', plugins:{ legend:{ position:'right', labels:{ color:'#8b95b0', font:{size:10}, padding:8 } } } }
    });

    // Source bar
    const srcK = Object.keys(stats.sources);
    charts.src = new Chart(document.getElementById('chart-src'), {
      type: 'bar',
      data: { labels: srcK, datasets:[{ data: srcK.map(k=>stats.sources[k]), backgroundColor: COLORS.slice(0,srcK.length), borderRadius:5, borderWidth:0 }] },
      options: { plugins:{legend:{display:false}}, scales:{ x:{ticks,grid}, y:{ticks,grid} } }
    });

    // Geo bar (horizontal)
    const geoCounts = {};
    ALL_JOBS.forEach(j => { const r = (j.location||'Remote · Worldwide').split('·')[0].trim(); geoCounts[r]=(geoCounts[r]||0)+1; });
    const geoK = Object.keys(geoCounts).sort((a,b)=>geoCounts[b]-geoCounts[a]).slice(0,8);
    charts.geo = new Chart(document.getElementById('chart-geo'), {
      type: 'bar',
      data: { labels: geoK, datasets:[{ data: geoK.map(k=>geoCounts[k]), backgroundColor:'#14b8a6', borderRadius:5, borderWidth:0 }] },
      options: { indexAxis:'y', plugins:{legend:{display:false}}, scales:{ x:{ticks,grid}, y:{ticks,grid:{display:false}} } }
    });

    // Time line
    charts.time = new Chart(document.getElementById('chart-time'), {
      type: 'line',
      data: {
        labels: stats.dates.map(d=>d.date),
        datasets:[{ data: stats.dates.map(d=>d.count), borderColor:'#6366f1', backgroundColor:'rgba(99,102,241,.1)', fill:true, tension:.4, pointRadius:3, pointBackgroundColor:'#6366f1', borderWidth:2 }]
      },
      options: { plugins:{legend:{display:false}}, scales:{ x:{ticks,grid}, y:{ticks,grid} } }
    });
  } catch(e) { toast('Could not load stats', '⚠️', 'toast-err'); }
}
function closeCharts() { document.getElementById('charts-modal').classList.remove('open'); }

// ── Toasts ────────────────────────────────────────────────────────────────────
function toast(msg, icon='', cls='') {
  const wrap = document.getElementById('toasts');
  const el   = document.createElement('div');
  el.className = `toast ${cls}`;
  el.innerHTML = `${icon?`<span>${icon}</span>`:''}<span>${esc(msg)}</span>`;
  wrap.appendChild(el);
  requestAnimationFrame(() => el.classList.add('show'));
  setTimeout(() => { el.classList.remove('show'); setTimeout(()=>el.remove(), 350); }, 2600);
}

// ── Utils ─────────────────────────────────────────────────────────────────────
function esc(s) {
  return String(s||'').replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;').replace(/"/g,'&quot;');
}
function timeAgo(iso) {
  if (!iso) return '';
  const ms  = Date.now() - new Date(iso).getTime();
  if (isNaN(ms)) return iso.slice(0,10);
  const m = Math.floor(ms/60000);
  if (m < 1)  return 'just now';
  if (m < 60) return `${m}m ago`;
  const h = Math.floor(m/60);
  if (h < 24) return `${h}h ago`;
  const d = Math.floor(h/24);
  if (d < 30) return `${d}d ago`;
  return new Date(iso).toLocaleDateString('en-US',{month:'short',day:'numeric'});
}
function show(id) { document.getElementById(id).style.display='block'; }
function hide(id) { document.getElementById(id).style.display='none'; }

// ── Card expand / collapse ───────────────────────────────────────────────────
document.addEventListener('click', e => {
  const card = e.target.closest('.job-card');
  if (!card) return;
  if (e.target.closest('a') || e.target.closest('button')) return;
  const expanded = card.classList.toggle('expanded');
  const hint = card.querySelector('.expand-hint');
  if (hint) hint.innerHTML = expanded
    ? `<svg class="expand-arrow" width="11" height="11" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5" style="transform:rotate(180deg)"><polyline points="6 9 12 15 18 9"/></svg> Less`
    : `<svg class="expand-arrow" width="11" height="11" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><polyline points="6 9 12 15 18 9"/></svg> Details`;
});

// ── Refresh ───────────────────────────────────────────────────────────────────
let _refreshPoll = null;

async function startRefresh() {
  const btn = document.getElementById('refresh-btn');
  if (btn.classList.contains('spinning')) { toast('Already refreshing…','⏳'); return; }
  try {
    const r = await fetch('/api/refresh', {method:'POST'});
    if (!r.ok) { const d=await r.json(); toast(d.error||'Could not start refresh','⚠️','toast-err'); return; }
    toast('Fetching new jobs in background…','🔄');
    btn.classList.add('spinning');
    document.getElementById('refresh-bar').classList.add('active');
    document.getElementById('refresh-fill').style.width = '0%';
    document.getElementById('nav-stat').innerHTML = '<b>0</b> / <b>?</b> feeds';
    _refreshPoll = setInterval(_pollRefresh, 1800);
  } catch(e) { toast('Network error','⚠️','toast-err'); }
}

async function _pollRefresh() {
  try {
    const r = await fetch('/api/refresh/status');
    const s = await r.json();
    const pct = s.total > 0 ? Math.round(s.done / s.total * 100) : 0;
    document.getElementById('refresh-fill').style.width = pct + '%';
    document.getElementById('nav-stat').innerHTML =
      `<b>${s.done}</b>/<b>${s.total}</b> feeds · <b>${s.added}</b> jobs`;
    if (!s.running) {
      clearInterval(_refreshPoll); _refreshPoll = null;
      document.getElementById('refresh-btn').classList.remove('spinning');
      document.getElementById('refresh-bar').classList.remove('active');
      toast(`Done — ${s.added} jobs fetched, ${s.failed} feeds failed`, '✅', 'toast-ok');
      await loadJobs();
      buildCheckLists();
      applyFilt();
    }
  } catch(e) {}
}

// ── Start ─────────────────────────────────────────────────────────────────────
boot();
</script>
</body>
</html>"""


# ── Frontend SPA ──────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return Response(DASHBOARD_HTML, mimetype="text/html; charset=utf-8")

# ── Applicant Profile Form ────────────────────────────────────────────────────

PROFILE_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Job Application Profile — JobAgent</title>
<style>
  :root{--bg:#0f1117;--card:#1a1d27;--border:#2a2d3a;--accent:#6c63ff;--accent2:#00d4aa;--text:#e2e8f0;--muted:#8892a4;--red:#f87171;--yellow:#fbbf24}
  *{box-sizing:border-box;margin:0;padding:0}
  body{background:var(--bg);color:var(--text);font-family:'Segoe UI',system-ui,sans-serif;font-size:14px;line-height:1.6;padding:24px 16px}
  h1{text-align:center;font-size:1.6rem;color:var(--accent);margin-bottom:4px}
  .subtitle{text-align:center;color:var(--muted);margin-bottom:32px;font-size:13px}
  .back{display:inline-flex;align-items:center;gap:6px;color:var(--muted);text-decoration:none;margin-bottom:20px;font-size:13px}
  .back:hover{color:var(--text)}
  .section{background:var(--card);border:1px solid var(--border);border-radius:10px;padding:24px;margin-bottom:20px}
  .section-title{font-size:1rem;font-weight:600;color:var(--accent2);margin-bottom:18px;display:flex;align-items:center;gap:8px}
  .section-title span{font-size:1.1rem}
  .grid{display:grid;gap:14px}
  .grid-2{grid-template-columns:1fr 1fr}
  .grid-3{grid-template-columns:1fr 1fr 1fr}
  @media(max-width:640px){.grid-2,.grid-3{grid-template-columns:1fr}}
  label{display:block;font-size:12px;color:var(--muted);margin-bottom:4px;font-weight:500}
  input,select,textarea{width:100%;background:#12151f;border:1px solid var(--border);border-radius:6px;color:var(--text);padding:9px 12px;font-size:13px;font-family:inherit;outline:none;transition:border-color .2s}
  input:focus,select:focus,textarea:focus{border-color:var(--accent)}
  textarea{resize:vertical;min-height:80px}
  select option{background:#12151f}
  .radio-group{display:flex;gap:16px;flex-wrap:wrap}
  .radio-group label{display:flex;align-items:center;gap:6px;color:var(--text);font-size:13px;font-weight:400;cursor:pointer}
  .radio-group input[type=radio]{width:auto;accent-color:var(--accent)}
  .checkbox-label{display:flex;align-items:center;gap:8px;cursor:pointer;color:var(--text);font-size:13px}
  .checkbox-label input{width:auto;accent-color:var(--accent)}
  .repeat-block{border:1px solid var(--border);border-radius:8px;padding:16px;margin-bottom:12px;position:relative}
  .repeat-block-title{font-size:12px;font-weight:600;color:var(--muted);margin-bottom:12px;text-transform:uppercase;letter-spacing:.05em}
  .add-btn{background:none;border:1px dashed var(--border);color:var(--muted);padding:8px 14px;border-radius:6px;cursor:pointer;font-size:13px;width:100%;margin-top:8px;transition:all .2s}
  .add-btn:hover{border-color:var(--accent);color:var(--accent)}
  .remove-btn{position:absolute;top:10px;right:10px;background:none;border:none;color:var(--muted);cursor:pointer;font-size:16px;line-height:1;padding:2px 6px;border-radius:4px}
  .remove-btn:hover{color:var(--red);background:#f8717120}
  .format-row{display:flex;gap:12px;align-items:center;flex-wrap:wrap;margin-bottom:18px}
  .format-btn{padding:10px 20px;border:1px solid var(--border);background:var(--card);color:var(--text);border-radius:6px;cursor:pointer;font-size:13px;transition:all .2s;display:flex;align-items:center;gap:6px}
  .format-btn.active{border-color:var(--accent);color:var(--accent);background:#6c63ff15}
  .format-btn:hover{border-color:var(--accent)}
  .submit-btn{width:100%;padding:14px;background:var(--accent);color:#fff;border:none;border-radius:8px;font-size:15px;font-weight:600;cursor:pointer;transition:opacity .2s;margin-top:8px}
  .submit-btn:hover{opacity:.88}
  .submit-btn:disabled{opacity:.4;cursor:default}
  .note{font-size:11px;color:var(--muted);margin-top:4px}
  .tag{background:#6c63ff20;color:var(--accent);border-radius:4px;padding:2px 7px;font-size:11px}
  #toast{position:fixed;bottom:24px;left:50%;transform:translateX(-50%);background:#1e2130;border:1px solid var(--border);padding:10px 20px;border-radius:8px;font-size:13px;display:none;z-index:999;min-width:220px;text-align:center}
  #toast.ok{border-color:var(--accent2);color:var(--accent2)}
  #toast.err{border-color:var(--red);color:var(--red)}
  .divider{border:none;border-top:1px solid var(--border);margin:14px 0}
</style>
</head>
<body>
<a href="/" class="back">← Back to Dashboard</a>
<h1>Job Application Profile</h1>
<p class="subtitle">Fill this out once — your browser agent will use it to apply automatically</p>

<form id="profileForm">

<!-- 1. Personal Info -->
<div class="section">
  <div class="section-title"><span>👤</span> Personal Information</div>
  <div class="grid grid-2">
    <div><label>First Name *</label><input name="first_name" required></div>
    <div><label>Last Name *</label><input name="last_name" required></div>
    <div><label>Email Address *</label><input name="email" type="email" required></div>
    <div><label>Phone Number *</label><input name="phone" type="tel" placeholder="+1 (555) 000-0000"></div>
    <div><label>City</label><input name="city" placeholder="e.g. Austin"></div>
    <div><label>State / Province</label><input name="state" placeholder="e.g. TX"></div>
    <div><label>Country</label><input name="country" value="United States"></div>
    <div><label>ZIP / Postal Code</label><input name="zip"></div>
  </div>
  <hr class="divider">
  <div class="grid grid-2">
    <div><label>LinkedIn URL</label><input name="linkedin" type="url" placeholder="https://linkedin.com/in/..."></div>
    <div><label>GitHub URL</label><input name="github" type="url" placeholder="https://github.com/..."></div>
    <div><label>Personal Website / Portfolio</label><input name="portfolio" type="url" placeholder="https://..."></div>
    <div><label>Other Profile (Dribbble, Behance, etc.)</label><input name="other_profile" type="url"></div>
  </div>
</div>

<!-- 2. Work Eligibility -->
<div class="section">
  <div class="section-title"><span>🛂</span> Work Eligibility & Logistics</div>
  <div class="grid">
    <div>
      <label>Are you legally authorized to work in the US?</label>
      <div class="radio-group">
        <label><input type="radio" name="work_authorized" value="Yes" checked> Yes</label>
        <label><input type="radio" name="work_authorized" value="No"> No</label>
      </div>
    </div>
    <div>
      <label>Do you now or will you in the future require visa sponsorship?</label>
      <div class="radio-group">
        <label><input type="radio" name="visa_sponsorship" value="No" checked> No</label>
        <label><input type="radio" name="visa_sponsorship" value="Yes, now"> Yes, now</label>
        <label><input type="radio" name="visa_sponsorship" value="Yes, in the future"> Yes, in the future</label>
      </div>
    </div>
    <div>
      <label>Visa / Work Authorization Type (if applicable)</label>
      <input name="visa_type" placeholder="e.g. US Citizen, Green Card, H-1B, OPT, etc.">
    </div>
    <div>
      <label>Work Preference</label>
      <div class="radio-group">
        <label><input type="radio" name="work_preference" value="Remote" checked> Remote</label>
        <label><input type="radio" name="work_preference" value="Hybrid"> Hybrid</label>
        <label><input type="radio" name="work_preference" value="On-site"> On-site</label>
        <label><input type="radio" name="work_preference" value="No preference"> No preference</label>
      </div>
    </div>
    <div>
      <label>Willing to Relocate?</label>
      <div class="radio-group">
        <label><input type="radio" name="relocate" value="Yes"> Yes</label>
        <label><input type="radio" name="relocate" value="No" checked> No</label>
        <label><input type="radio" name="relocate" value="Maybe"> Maybe / Open to it</label>
      </div>
    </div>
    <div class="grid-2 grid">
      <div><label>Available Start Date</label><input name="start_date" type="date"></div>
      <div><label>Notice Period (at current job)</label><input name="notice_period" placeholder="e.g. 2 weeks, 1 month, Immediately"></div>
    </div>
    <div>
      <label>Employment Type</label>
      <div class="radio-group">
        <label><input type="radio" name="employment_type" value="Full-time" checked> Full-time</label>
        <label><input type="radio" name="employment_type" value="Part-time"> Part-time</label>
        <label><input type="radio" name="employment_type" value="Contract"> Contract</label>
        <label><input type="radio" name="employment_type" value="Open to all"> Open to all</label>
      </div>
    </div>
  </div>
</div>

<!-- 3. Work History -->
<div class="section">
  <div class="section-title"><span>💼</span> Work History</div>
  <p class="note" style="margin-bottom:14px">Add your most recent positions first. Many applications require re-entering resume details.</p>
  <div id="jobs-container">
    <div class="repeat-block" data-idx="0">
      <div class="repeat-block-title">Position 1 <span class="tag">Current / Most Recent</span></div>
      <button type="button" class="remove-btn" onclick="removeBlock(this,'jobs-container')">×</button>
      <div class="grid grid-2">
        <div><label>Job Title</label><input name="job_title_0"></div>
        <div><label>Company Name</label><input name="company_0"></div>
        <div><label>Start Date</label><input name="job_start_0" type="month"></div>
        <div><label>End Date</label><input name="job_end_0" type="month" placeholder="Leave blank if current"></div>
        <div><label>City / Location</label><input name="job_location_0"></div>
        <div><label>Employment Type</label>
          <select name="job_type_0">
            <option>Full-time</option><option>Part-time</option><option>Contract</option><option>Internship</option><option>Freelance</option>
          </select>
        </div>
      </div>
      <div style="margin-top:12px"><label>Key Responsibilities & Achievements</label><textarea name="job_desc_0" rows="3" placeholder="Bullet points work great here"></textarea></div>
      <div style="margin-top:10px"><label>Manager Name (if asked)</label><input name="job_manager_0"></div>
    </div>
  </div>
  <button type="button" class="add-btn" onclick="addJob()">+ Add Another Position</button>
</div>

<!-- 4. Education -->
<div class="section">
  <div class="section-title"><span>🎓</span> Education</div>
  <div id="edu-container">
    <div class="repeat-block" data-idx="0">
      <div class="repeat-block-title">Degree 1</div>
      <button type="button" class="remove-btn" onclick="removeBlock(this,'edu-container')">×</button>
      <div class="grid grid-2">
        <div><label>School / University</label><input name="school_0"></div>
        <div><label>Degree (e.g. B.S., M.S., Associate)</label><input name="degree_0"></div>
        <div><label>Major / Field of Study</label><input name="major_0"></div>
        <div><label>GPA (optional)</label><input name="gpa_0" placeholder="e.g. 3.7"></div>
        <div><label>Graduation Date</label><input name="grad_date_0" type="month"></div>
        <div><label>Location (City, State)</label><input name="school_location_0"></div>
      </div>
    </div>
  </div>
  <button type="button" class="add-btn" onclick="addEdu()">+ Add Another Degree / School</button>
</div>

<!-- 5. Skills & Certifications -->
<div class="section">
  <div class="section-title"><span>🛠️</span> Skills & Certifications</div>
  <div class="grid">
    <div><label>Technical Skills (comma-separated)</label><input name="tech_skills" placeholder="e.g. Python, SQL, React, AWS, Docker"></div>
    <div><label>Tools & Software</label><input name="tools" placeholder="e.g. Jira, Figma, Salesforce, Excel"></div>
    <div><label>Languages Spoken</label><input name="languages" placeholder="e.g. English (Native), Spanish (Intermediate)"></div>
    <div id="certs-container">
      <div class="repeat-block" data-idx="0">
        <div class="repeat-block-title">Certification 1</div>
        <button type="button" class="remove-btn" onclick="removeBlock(this,'certs-container')">×</button>
        <div class="grid grid-3">
          <div><label>Certification Name</label><input name="cert_name_0"></div>
          <div><label>Issuing Organization</label><input name="cert_org_0"></div>
          <div><label>Date Earned</label><input name="cert_date_0" type="month"></div>
        </div>
      </div>
    </div>
    <button type="button" class="add-btn" onclick="addCert()">+ Add Certification / License</button>
  </div>
</div>

<!-- 6. Screening Questions -->
<div class="section">
  <div class="section-title"><span>🎯</span> Common Screening Questions</div>
  <p class="note" style="margin-bottom:14px">Pre-written answers your agent can use or adapt.</p>
  <div class="grid">
    <div>
      <label>Why do you want to work at [Company]? (Template — agent will customize)</label>
      <textarea name="why_company" rows="3" placeholder="I'm excited about [Company] because..."></textarea>
    </div>
    <div>
      <label>Describe a relevant project or achievement</label>
      <textarea name="achievement" rows="3" placeholder="In my previous role at X, I led a project that..."></textarea>
    </div>
    <div>
      <label>What are your greatest strengths?</label>
      <textarea name="strengths" rows="2"></textarea>
    </div>
    <div>
      <label>How do you handle a challenging situation / conflict?</label>
      <textarea name="challenge_answer" rows="2"></textarea>
    </div>
    <div>
      <label>Where do you see yourself in 5 years?</label>
      <textarea name="five_years" rows="2"></textarea>
    </div>
    <div>
      <label>Any additional notes for your agent (custom instructions)</label>
      <textarea name="agent_notes" rows="2" placeholder="e.g. Always mention my open-source work, never disclose current salary"></textarea>
    </div>
  </div>
</div>

<!-- 7. Portfolio / Work Samples -->
<div class="section">
  <div class="section-title"><span>🔗</span> Portfolio & Work Samples</div>
  <div id="portfolio-container">
    <div class="repeat-block" data-idx="0">
      <div class="repeat-block-title">Sample 1</div>
      <button type="button" class="remove-btn" onclick="removeBlock(this,'portfolio-container')">×</button>
      <div class="grid grid-2">
        <div><label>Title / Project Name</label><input name="sample_title_0"></div>
        <div><label>URL</label><input name="sample_url_0" type="url"></div>
        <div style="grid-column:1/-1"><label>Brief Description</label><input name="sample_desc_0"></div>
      </div>
    </div>
  </div>
  <button type="button" class="add-btn" onclick="addSample()">+ Add Work Sample / Project</button>
</div>

<!-- 8. References -->
<div class="section">
  <div class="section-title"><span>📋</span> References</div>
  <p class="note" style="margin-bottom:14px">Usually requested after initial application, but helpful to have ready.</p>
  <div id="refs-container">
    <div class="repeat-block" data-idx="0">
      <div class="repeat-block-title">Reference 1</div>
      <button type="button" class="remove-btn" onclick="removeBlock(this,'refs-container')">×</button>
      <div class="grid grid-2">
        <div><label>Name</label><input name="ref_name_0"></div>
        <div><label>Title / Role</label><input name="ref_title_0"></div>
        <div><label>Company</label><input name="ref_company_0"></div>
        <div><label>Relationship to You</label><input name="ref_relationship_0" placeholder="e.g. Former Manager"></div>
        <div><label>Email</label><input name="ref_email_0" type="email"></div>
        <div><label>Phone</label><input name="ref_phone_0" type="tel"></div>
      </div>
    </div>
  </div>
  <button type="button" class="add-btn" onclick="addRef()">+ Add Reference</button>
</div>

<!-- 9. Compensation -->
<div class="section">
  <div class="section-title"><span>💰</span> Compensation Expectations</div>
  <div class="grid grid-2">
    <div><label>Desired Salary (Minimum)</label><input name="salary_min" placeholder="e.g. 90000"></div>
    <div><label>Desired Salary (Maximum / Ideal)</label><input name="salary_max" placeholder="e.g. 120000"></div>
    <div><label>Currency</label>
      <select name="salary_currency">
        <option>USD</option><option>CAD</option><option>GBP</option><option>EUR</option><option>AUD</option>
      </select>
    </div>
    <div><label>Pay Frequency</label>
      <select name="pay_frequency">
        <option>Annual</option><option>Hourly</option><option>Monthly</option>
      </select>
    </div>
    <div style="grid-column:1/-1"><label>Notes on Compensation (what to say if asked)</label><input name="salary_note" placeholder="e.g. Negotiable based on total comp, open to equity discussion"></div>
  </div>
</div>

<!-- 10. EEO & Background -->
<div class="section">
  <div class="section-title"><span>📝</span> EEO & Compliance <span class="tag" style="font-size:11px">Optional — for compliance forms only</span></div>
  <div class="grid grid-2">
    <div>
      <label>Gender</label>
      <select name="gender">
        <option value="">Prefer not to say</option>
        <option>Male</option><option>Female</option><option>Non-binary</option><option>Self-describe</option>
      </select>
    </div>
    <div>
      <label>Race / Ethnicity</label>
      <select name="ethnicity">
        <option value="">Prefer not to say</option>
        <option>Hispanic or Latino</option>
        <option>White (not Hispanic)</option>
        <option>Black or African American</option>
        <option>Asian</option>
        <option>Native Hawaiian or Other Pacific Islander</option>
        <option>American Indian or Alaskan Native</option>
        <option>Two or more races</option>
      </select>
    </div>
    <div>
      <label>Veteran Status</label>
      <select name="veteran_status">
        <option value="">Prefer not to say</option>
        <option>Not a veteran</option>
        <option>Protected veteran</option>
        <option>Active duty</option>
      </select>
    </div>
    <div>
      <label>Disability Status</label>
      <select name="disability_status">
        <option value="">Prefer not to say</option>
        <option>No disability</option>
        <option>Yes, I have a disability</option>
      </select>
    </div>
  </div>
  <hr class="divider">
  <div class="grid">
    <div class="checkbox-label">
      <input type="checkbox" name="background_check_consent" value="yes">
      <span>I consent to a background check as part of the application process</span>
    </div>
    <div class="checkbox-label">
      <input type="checkbox" name="drug_screen_consent" value="yes">
      <input type="hidden" name="drug_screen_consent_flag" value="no">
      <span>I consent to drug screening (for roles that require it)</span>
    </div>
  </div>
</div>

<!-- Export format + Submit -->
<div class="section">
  <div class="section-title"><span>💾</span> Save Profile As</div>
  <div class="format-row">
    <button type="button" class="format-btn active" id="btn-docx" onclick="setFormat('docx')">📄 Word (.docx)</button>
    <button type="button" class="format-btn" id="btn-pdf" onclick="setFormat('pdf')">📋 PDF</button>
    <button type="button" class="format-btn" id="btn-json" onclick="setFormat('json')">⚙️ JSON (for agents)</button>
  </div>
  <input type="hidden" name="export_format" id="export_format" value="docx">
  <button type="submit" class="submit-btn" id="submitBtn">Save My Profile</button>
</div>

</form>

<div id="toast"></div>

<script>
let jobCount = 1, eduCount = 1, certCount = 1, sampleCount = 1, refCount = 1;

function setFormat(fmt) {
  document.getElementById('export_format').value = fmt;
  ['docx','pdf','json'].forEach(f => {
    document.getElementById('btn-'+f).classList.toggle('active', f===fmt);
  });
}

function removeBlock(btn, containerId) {
  const container = document.getElementById(containerId);
  if (container.querySelectorAll('.repeat-block').length <= 1) return;
  btn.closest('.repeat-block').remove();
}

function addJob() {
  const c = document.getElementById('jobs-container');
  const d = document.createElement('div');
  d.className = 'repeat-block'; d.dataset.idx = jobCount;
  d.innerHTML = `<div class="repeat-block-title">Position ${c.querySelectorAll('.repeat-block').length + 1}</div>
    <button type="button" class="remove-btn" onclick="removeBlock(this,'jobs-container')">×</button>
    <div class="grid grid-2">
      <div><label>Job Title</label><input name="job_title_${jobCount}"></div>
      <div><label>Company Name</label><input name="company_${jobCount}"></div>
      <div><label>Start Date</label><input name="job_start_${jobCount}" type="month"></div>
      <div><label>End Date</label><input name="job_end_${jobCount}" type="month"></div>
      <div><label>City / Location</label><input name="job_location_${jobCount}"></div>
      <div><label>Employment Type</label>
        <select name="job_type_${jobCount}"><option>Full-time</option><option>Part-time</option><option>Contract</option><option>Internship</option><option>Freelance</option></select>
      </div>
    </div>
    <div style="margin-top:12px"><label>Key Responsibilities & Achievements</label><textarea name="job_desc_${jobCount}" rows="3"></textarea></div>
    <div style="margin-top:10px"><label>Manager Name</label><input name="job_manager_${jobCount}"></div>`;
  c.appendChild(d); jobCount++;
}

function addEdu() {
  const c = document.getElementById('edu-container');
  const d = document.createElement('div');
  d.className = 'repeat-block'; d.dataset.idx = eduCount;
  d.innerHTML = `<div class="repeat-block-title">Degree ${c.querySelectorAll('.repeat-block').length + 1}</div>
    <button type="button" class="remove-btn" onclick="removeBlock(this,'edu-container')">×</button>
    <div class="grid grid-2">
      <div><label>School / University</label><input name="school_${eduCount}"></div>
      <div><label>Degree</label><input name="degree_${eduCount}"></div>
      <div><label>Major</label><input name="major_${eduCount}"></div>
      <div><label>GPA</label><input name="gpa_${eduCount}"></div>
      <div><label>Graduation Date</label><input name="grad_date_${eduCount}" type="month"></div>
      <div><label>Location</label><input name="school_location_${eduCount}"></div>
    </div>`;
  c.appendChild(d); eduCount++;
}

function addCert() {
  const c = document.getElementById('certs-container');
  const d = document.createElement('div');
  d.className = 'repeat-block'; d.dataset.idx = certCount;
  d.innerHTML = `<div class="repeat-block-title">Certification ${c.querySelectorAll('.repeat-block').length + 1}</div>
    <button type="button" class="remove-btn" onclick="removeBlock(this,'certs-container')">×</button>
    <div class="grid grid-3">
      <div><label>Certification Name</label><input name="cert_name_${certCount}"></div>
      <div><label>Issuing Organization</label><input name="cert_org_${certCount}"></div>
      <div><label>Date Earned</label><input name="cert_date_${certCount}" type="month"></div>
    </div>`;
  c.appendChild(d); certCount++;
}

function addSample() {
  const c = document.getElementById('portfolio-container');
  const d = document.createElement('div');
  d.className = 'repeat-block'; d.dataset.idx = sampleCount;
  d.innerHTML = `<div class="repeat-block-title">Sample ${c.querySelectorAll('.repeat-block').length + 1}</div>
    <button type="button" class="remove-btn" onclick="removeBlock(this,'portfolio-container')">×</button>
    <div class="grid grid-2">
      <div><label>Title / Project Name</label><input name="sample_title_${sampleCount}"></div>
      <div><label>URL</label><input name="sample_url_${sampleCount}" type="url"></div>
      <div style="grid-column:1/-1"><label>Brief Description</label><input name="sample_desc_${sampleCount}"></div>
    </div>`;
  c.appendChild(d); sampleCount++;
}

function addRef() {
  const c = document.getElementById('refs-container');
  const d = document.createElement('div');
  d.className = 'repeat-block'; d.dataset.idx = refCount;
  d.innerHTML = `<div class="repeat-block-title">Reference ${c.querySelectorAll('.repeat-block').length + 1}</div>
    <button type="button" class="remove-btn" onclick="removeBlock(this,'refs-container')">×</button>
    <div class="grid grid-2">
      <div><label>Name</label><input name="ref_name_${refCount}"></div>
      <div><label>Title / Role</label><input name="ref_title_${refCount}"></div>
      <div><label>Company</label><input name="ref_company_${refCount}"></div>
      <div><label>Relationship</label><input name="ref_relationship_${refCount}"></div>
      <div><label>Email</label><input name="ref_email_${refCount}" type="email"></div>
      <div><label>Phone</label><input name="ref_phone_${refCount}" type="tel"></div>
    </div>`;
  c.appendChild(d); refCount++;
}

function toast(msg, type='ok') {
  const t = document.getElementById('toast');
  t.textContent = msg; t.className = type; t.style.display = 'block';
  setTimeout(() => t.style.display = 'none', 3500);
}

document.getElementById('profileForm').addEventListener('submit', async function(e) {
  e.preventDefault();
  const btn = document.getElementById('submitBtn');
  btn.disabled = true; btn.textContent = 'Saving…';
  const fd = new FormData(this);
  const fmt = document.getElementById('export_format').value;
  try {
    const res = await fetch('/api/profile/save', {method:'POST', body: fd});
    if (!res.ok) {
      const err = await res.json();
      toast(err.error || 'Save failed', 'err');
    } else {
      const blob = await res.blob();
      const cd = res.headers.get('Content-Disposition') || '';
      const fname = (cd.match(/filename="([^"]+)"/) || [])[1] || 'job_profile.' + fmt;
      const url = URL.createObjectURL(blob);
      const a = document.createElement('a');
      a.href = url; a.download = fname; a.click();
      URL.revokeObjectURL(url);
      toast('Profile saved and downloaded!', 'ok');
    }
  } catch(err) {
    toast('Network error', 'err');
  }
  btn.disabled = false;
  btn.textContent = 'Save My Profile';
});
</script>
</body>
</html>"""


def _collect_repeating(data, prefix_map):
    """Extract indexed repeating blocks from flat form data into list of dicts."""
    results = []
    idx = 0
    while True:
        first_key = list(prefix_map.keys())[0]
        if f"{first_key}{idx}" not in data and idx > 0:
            break
        entry = {}
        for field, label in prefix_map.items():
            entry[label] = data.get(f"{field}{idx}", "").strip()
        if any(entry.values()):
            results.append(entry)
        idx += 1
        if idx > 20:
            break
    return results


def _add_section(doc, title):
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4C, 0x5F, 0xFF)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph("─" * 60).paragraph_format.space_after = Pt(4)


def _add_field(doc, label, value):
    if not value:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.size = Pt(10)
    p.add_run(str(value)).font.size = Pt(10)


def build_docx(data):
    doc = Document()
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("JOB APPLICATION PROFILE")
    run.bold = True
    run.font.size = Pt(16)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Generated {datetime.now().strftime('%B %d, %Y')}").font.size = Pt(9)

    _add_section(doc, "Personal Information")
    name = f"{data.get('first_name','')} {data.get('last_name','')}".strip()
    _add_field(doc, "Full Name", name)
    for f, l in [("email","Email"),("phone","Phone"),("linkedin","LinkedIn"),
                 ("github","GitHub"),("portfolio","Portfolio"),("other_profile","Other Profile")]:
        _add_field(doc, l, data.get(f,""))
    loc = ", ".join(filter(None, [data.get("city",""), data.get("state",""), data.get("country",""), data.get("zip","")]))
    _add_field(doc, "Location", loc)

    _add_section(doc, "Work Eligibility & Logistics")
    for f, l in [("work_authorized","Authorized to Work in US"),("visa_sponsorship","Visa Sponsorship Needed"),
                 ("visa_type","Visa / Auth Type"),("work_preference","Work Preference"),("relocate","Willing to Relocate"),
                 ("start_date","Available Start Date"),("notice_period","Notice Period"),("employment_type","Employment Type")]:
        _add_field(doc, l, data.get(f,""))

    _add_section(doc, "Work History")
    jobs = _collect_repeating(data, {"job_title_":"Title","company_":"Company","job_start_":"Start",
                                      "job_end_":"End","job_location_":"Location","job_type_":"Type",
                                      "job_desc_":"Responsibilities","job_manager_":"Manager"})
    for i, j in enumerate(jobs, 1):
        doc.add_paragraph(f"Position {i}:").runs[0].bold = True
        for k, v in j.items():
            _add_field(doc, k, v)
        doc.add_paragraph("")

    _add_section(doc, "Education")
    edu = _collect_repeating(data, {"school_":"School","degree_":"Degree","major_":"Major",
                                     "gpa_":"GPA","grad_date_":"Graduation","school_location_":"Location"})
    for i, e in enumerate(edu, 1):
        doc.add_paragraph(f"Education {i}:").runs[0].bold = True
        for k, v in e.items():
            _add_field(doc, k, v)
        doc.add_paragraph("")

    _add_section(doc, "Skills & Certifications")
    _add_field(doc, "Technical Skills", data.get("tech_skills",""))
    _add_field(doc, "Tools & Software", data.get("tools",""))
    _add_field(doc, "Languages", data.get("languages",""))
    certs = _collect_repeating(data, {"cert_name_":"Name","cert_org_":"Issuer","cert_date_":"Date"})
    for i, c in enumerate(certs, 1):
        if any(c.values()):
            doc.add_paragraph(f"Cert {i}: {c.get('Name','')} — {c.get('Issuer','')} ({c.get('Date','')})").font = None

    _add_section(doc, "Screening Questions")
    for f, l in [("why_company","Why This Company"),("achievement","Key Achievement"),
                 ("strengths","Strengths"),("challenge_answer","Handling Challenges"),
                 ("five_years","5-Year Vision"),("agent_notes","Agent Instructions")]:
        _add_field(doc, l, data.get(f,""))

    _add_section(doc, "Portfolio & Work Samples")
    samples = _collect_repeating(data, {"sample_title_":"Title","sample_url_":"URL","sample_desc_":"Description"})
    for i, s in enumerate(samples, 1):
        if any(s.values()):
            doc.add_paragraph(f"Sample {i}: {s.get('Title','')} — {s.get('URL','')}").runs[0].font.size = Pt(10)
            if s.get("Description"):
                doc.add_paragraph(f"  {s['Description']}").runs[0].font.size = Pt(10)

    _add_section(doc, "References")
    refs = _collect_repeating(data, {"ref_name_":"Name","ref_title_":"Title","ref_company_":"Company",
                                      "ref_relationship_":"Relationship","ref_email_":"Email","ref_phone_":"Phone"})
    for i, r in enumerate(refs, 1):
        doc.add_paragraph(f"Reference {i}:").runs[0].bold = True
        for k, v in r.items():
            _add_field(doc, k, v)
        doc.add_paragraph("")

    _add_section(doc, "Compensation")
    low = data.get("salary_min",""); high = data.get("salary_max","")
    cur = data.get("salary_currency","USD"); freq = data.get("pay_frequency","Annual")
    if low or high:
        _add_field(doc, "Desired Range", f"{cur} {low} – {high} ({freq})")
    _add_field(doc, "Compensation Notes", data.get("salary_note",""))

    _add_section(doc, "EEO & Compliance")
    for f, l in [("gender","Gender"),("ethnicity","Ethnicity"),("veteran_status","Veteran Status"),
                 ("disability_status","Disability Status"),("background_check_consent","Background Check Consent"),
                 ("drug_screen_consent","Drug Screen Consent")]:
        _add_field(doc, l, data.get(f,""))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.route("/profile")
def profile_page():
    return Response(PROFILE_HTML, mimetype="text/html; charset=utf-8")


@app.route("/api/profile/save", methods=["POST"])
def save_profile():
    data = request.form.to_dict()
    fmt = data.get("export_format", "docx").lower()

    if fmt == "json":
        fname = f"job_profile_{datetime.now().strftime('%Y%m%d')}.json"
        return Response(
            json.dumps(data, indent=2),
            mimetype="application/json",
            headers={"Content-Disposition": f'attachment; filename="{fname}"'}
        )

    if fmt == "pdf":
        try:
            from docx2pdf import convert
            docx_buf = build_docx(data)
            tmp_docx = Path("/tmp/job_profile_tmp.docx")
            tmp_pdf  = Path("/tmp/job_profile_tmp.pdf")
            tmp_docx.write_bytes(docx_buf.read())
            convert(str(tmp_docx), str(tmp_pdf))
            pdf_bytes = tmp_pdf.read_bytes()
            fname = f"job_profile_{datetime.now().strftime('%Y%m%d')}.pdf"
            return Response(
                pdf_bytes,
                mimetype="application/pdf",
                headers={"Content-Disposition": f'attachment; filename="{fname}"'}
            )
        except ImportError:
            return jsonify({"error": "PDF export requires docx2pdf. Install it with: pip install docx2pdf — or use Word (.docx) format instead."}), 400
        except Exception as e:
            return jsonify({"error": f"PDF conversion failed: {e}"}), 500

    # Default: docx
    buf = build_docx(data)
    fname = f"job_profile_{datetime.now().strftime('%Y%m%d')}.docx"
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=fname
    )


if __name__ == "__main__":
    print("\n  JobAgent  →  http://localhost:5000\n")
    app.run(debug=False, port=5000)
